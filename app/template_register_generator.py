from copy import deepcopy
from math import ceil
from pathlib import Path
import shutil

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from app.attendance_register_generator import get_students_for_group
from app.config import DOCX_TEMPLATES_DIR
from app.docx_utils import ensure_parent, format_time_range, safe_filename_text
from app.template_claim_generator import MissingDocxTemplateError, TemplateMappingError, require_template


ATTENDANCE_TEMPLATE_PATH = DOCX_TEMPLATES_DIR / "golden_attendance_register_template.docx"


def _set_paragraph_text(paragraph, text: str, font_size: int | None = None, alignment=None) -> None:
    if not paragraph.runs:
        run = paragraph.add_run(str(text))
        return
    paragraph.runs[0].text = str(text)
    for run in paragraph.runs[1:]:
        run.text = ""


def _set_cell_text_preserve(cell, text: str, font_size: int | None = None, alignment=None) -> None:
    paragraph = cell.paragraphs[0]
    _set_paragraph_text(paragraph, str(text), font_size=font_size, alignment=alignment)
    for extra in cell.paragraphs[1:]:
        _set_paragraph_text(extra, "")


def _format_register_date(value) -> str:
    return pd.to_datetime(value).strftime("%d-%m-%y")


def _compact_department(value: str) -> str:
    text = str(value).strip()
    if "Informatics" in text:
        return "Informatics"
    return text


def map_register_template(document: Document) -> dict:
    if not document.tables:
        raise TemplateMappingError("Attendance register template does not contain a table")
    table = document.tables[0]
    date_row = time_row = None
    for index, row in enumerate(table.rows):
        row_text = " ".join(cell.text for cell in row.cells).lower()
        if "date" in row_text and date_row is None:
            date_row = index
        if "time" in row_text and time_row is None:
            time_row = index
    if date_row is None or time_row is None:
        raise TemplateMappingError("Could not identify DATE/TIME rows in attendance template")

    student_start = time_row + 1
    student_end = len(table.rows) - 1
    session_columns = list(range(4, max((len(row.cells) for row in table.rows), default=9)))
    return {
        "table_index": 0,
        "date_row": date_row,
        "time_row": time_row,
        "student_start_row": student_start,
        "student_end_row": student_end,
        "number_column": 0,
        "student_surname_column": 1,
        "student_initials_column": 2,
        "student_number_column": 3,
        "session_columns": session_columns,
        "footer_paragraph_range": (7, 8),
        "paragraphs": {
            "faculty": 1,
            "department": 2,
            "course": 3,
            "signature_values": 7,
            "signature_labels": 8,
        },
    }


def _clone_template_page(output: Document, template: Document) -> int:
    output.add_page_break()
    start_index = len(output.paragraphs)
    body = output._body._element
    for child in template._body._element:
        if child.tag.endswith("}sectPr"):
            continue
        body.append(deepcopy(child))
    return start_index


def _append_row_like(table, template_row_index: int):
    new_tr = deepcopy(table.rows[template_row_index]._tr)
    table._tbl.append(new_tr)
    return table.rows[-1]


def _delete_row(table, row_index: int) -> None:
    row = table.rows[row_index]
    table._tbl.remove(row._tr)


def _fill_paragraphs(
    paragraphs,
    mapping: dict,
    faculty: str,
    department: str,
    course_name: str,
    course_code: str,
    group_name: str,
    lecturer_name: str,
    staff_number: str,
) -> None:
    idx = mapping["paragraphs"]
    _set_paragraph_text(paragraphs[idx["faculty"]], f"FACULTY:\t___{faculty}______________________________________________________________________")
    _set_paragraph_text(paragraphs[idx["department"]], f"DEPARTMENT: {_compact_department(department)}", font_size=9)
    _set_paragraph_text(
        paragraphs[idx["course"]],
        f"COURSE NAME: {course_name}\t                     COURSE CODE: {course_code}                                                                 GROUP:  {group_name}",
    )
    _set_paragraph_text(
        paragraphs[idx["signature_values"]],
        f"{lecturer_name}\t _______________________________\t\t{staff_number}\t",
        font_size=9,
    )


def _student_rows(students: list[dict]) -> list[tuple[str, str, str]]:
    rows = []
    for student in students:
        surname = str(student.get("surname") or "").strip()
        initials = str(student.get("initials") or "").strip()
        if not surname and student.get("full_name"):
            parts = str(student["full_name"]).split()
            surname = parts[-1] if parts else ""
            initials = "".join(part[0] for part in parts[:-1] if part)
        rows.append((surname, initials, str(student["student_number"])))
    return rows


def _fill_table(table, mapping: dict, students: list[dict], session_chunk: list[dict]) -> None:
    date_row = table.rows[mapping["date_row"]]
    time_row = table.rows[mapping["time_row"]]
    session_columns = mapping["session_columns"]

    for column in session_columns:
        _set_cell_text_preserve(date_row.cells[column], "", font_size=8, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell_text_preserve(time_row.cells[column], "", font_size=8, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    for column, session in zip(session_columns, session_chunk):
        _set_cell_text_preserve(
            date_row.cells[column],
            _format_register_date(session["session_date"]),
            font_size=8,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
        )
        _set_cell_text_preserve(
            time_row.cells[column],
            format_time_range(session["start_time"], session["end_time"]),
            font_size=8,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
        )

    student_values = _student_rows(students)
    available = mapping["student_end_row"] - mapping["student_start_row"] + 1
    if len(student_values) > available:
        source = mapping["student_start_row"]
        for _ in range(len(student_values) - available):
            _append_row_like(table, source)
        mapping["student_end_row"] = len(table.rows) - 1

    for row_index in range(mapping["student_start_row"], mapping["student_end_row"] + 1):
        row = table.rows[row_index]
        _set_cell_text_preserve(
            row.cells[mapping["number_column"]],
            row.cells[mapping["number_column"]].text,
            font_size=9,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
        )
        _set_cell_text_preserve(
            row.cells[mapping["student_surname_column"]],
            "",
            font_size=9,
            alignment=WD_ALIGN_PARAGRAPH.LEFT,
        )
        _set_cell_text_preserve(
            row.cells[mapping["student_initials_column"]],
            "",
            font_size=9,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
        )
        _set_cell_text_preserve(row.cells[mapping["student_number_column"]], "", font_size=8, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        for column in session_columns:
            _set_cell_text_preserve(row.cells[column], "", font_size=8, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    for offset, (surname, initials, student_number) in enumerate(student_values):
        row = table.rows[mapping["student_start_row"] + offset]
        _set_cell_text_preserve(
            row.cells[mapping["number_column"]],
            row.cells[mapping["number_column"]].text,
            font_size=9,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
        )
        _set_cell_text_preserve(
            row.cells[mapping["student_surname_column"]],
            surname,
            font_size=9,
            alignment=WD_ALIGN_PARAGRAPH.LEFT,
        )
        _set_cell_text_preserve(
            row.cells[mapping["student_initials_column"]],
            initials,
            font_size=9,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
        )
        _set_cell_text_preserve(row.cells[mapping["student_number_column"]], student_number, font_size=8, alignment=WD_ALIGN_PARAGRAPH.CENTER)


def generate_template_attendance_register_pack(
    sessions_df: pd.DataFrame,
    output_path: str | Path,
    year: int,
    month: int,
    warning: bool = False,
    template_path: str | Path = ATTENDANCE_TEMPLATE_PATH,
    strict: bool = True,
) -> list[Path]:
    if sessions_df.empty:
        raise ValueError("Cannot generate attendance register pack without sessions")

    template_file = require_template(template_path)
    output = ensure_parent(output_path)
    output_dir = output.parent / "registers"
    output_dir.mkdir(parents=True, exist_ok=True)
    template_document = Document(template_file)
    base_mapping = map_register_template(template_document)
    max_session_columns = len(base_mapping["session_columns"])

    lecturer_name = str(sessions_df["lecturer_name"].iloc[0])
    staff_number = str(sessions_df["staff_number"].iloc[0])
    generated_paths: list[Path] = []

    group_keys = ["faculty", "department", "course_name", "course_code", "group_name", "campus"]
    for group_values, group_df in sessions_df.groupby(group_keys, sort=True):
        faculty, department, course_name, course_code, group_name, _campus = group_values
        group_df = group_df.sort_values(["session_date", "start_time"])
        students = get_students_for_group(group_name, course_code)
        page_count = ceil(len(group_df) / max_session_columns)
        for chunk_index in range(page_count):
            file_name = (
                f"register_{staff_number}_{safe_filename_text(str(course_code))}_"
                f"{safe_filename_text(str(group_name))}_{year}_{month:02d}_p{chunk_index + 1}.docx"
            )
            page_output = output_dir / file_name
            shutil.copy2(template_file, page_output)
            output_document = Document(page_output)
            page_mapping = deepcopy(base_mapping)
            paragraphs = output_document.paragraphs[: len(template_document.paragraphs)]
            table = output_document.tables[0]
            chunk = group_df.iloc[
                chunk_index * max_session_columns : (chunk_index + 1) * max_session_columns
            ].to_dict("records")
            _fill_paragraphs(
                paragraphs,
                page_mapping,
                faculty,
                department,
                course_name,
                course_code,
                group_name,
                lecturer_name,
                staff_number,
            )
            _fill_table(table, page_mapping, students, chunk)

            output_document.save(page_output)
            generated_paths.append(page_output)

    return generated_paths
