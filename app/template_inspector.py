import argparse
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

from app.template_claim_generator import map_claim_template
from app.template_register_generator import map_register_template


KEY_LABELS = [
    "Name & Surname",
    "Highest qualification",
    "Personnel Number",
    "Budget Allocation",
    "Tariff per hour",
    "Identity / Passport Number",
    "PAYE No.",
    "Address",
    "Tel. no.",
    "PARTICULARS OF CLAIM",
    "Claimant's Signature",
    "Claimant’s Signature",
    "CLASS ATTENDANCE SHEET",
    "FACULTY",
    "DEPARTMENT",
    "COURSE NAME",
    "COURSE CODE",
    "GROUP",
    "DATE",
    "TIME",
    "NAME OF LECTURER",
    "SIGNATURE",
    "STAFF NR.",
]


def document_text(document: Document) -> str:
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


def inspect_template(path: str | Path) -> dict:
    template_path = Path(path)
    if not template_path.exists():
        raise FileNotFoundError(f"Template not found: {template_path}")

    document = Document(template_path)
    full_text = document_text(document)
    tables = []
    for index, table in enumerate(document.tables):
        row_count = len(table.rows)
        col_count = max((len(row.cells) for row in table.rows), default=0)
        sample_cells: list[str] = []
        for row in table.rows[:3]:
            sample_cells.extend(cell.text.replace("\n", " / ").strip() for cell in row.cells[:6])
        tables.append(
            {
                "index": index,
                "rows": row_count,
                "columns": col_count,
                "sample_cells": sample_cells[:12],
            }
        )

    found_labels = [label for label in KEY_LABELS if label.lower() in full_text.lower()]
    return {
        "path": str(template_path),
        "paragraph_count": len(document.paragraphs),
        "table_count": len(document.tables),
        "tables": tables,
        "labels_found": found_labels,
    }


def print_template_analysis(analysis: dict) -> None:
    print(f"Document path: {analysis['path']}")
    print(f"Number of paragraphs: {analysis['paragraph_count']}")
    print(f"Number of tables: {analysis['table_count']}")
    print("Table dimensions and sample cells:")
    for table in analysis["tables"]:
        print(f"- Table {table['index']}: {table['rows']} rows x {table['columns']} columns")
        if table["sample_cells"]:
            print(f"  First cells: {' | '.join(table['sample_cells'])}")
    print("Key labels detected:")
    if analysis["labels_found"]:
        for label in analysis["labels_found"]:
            print(f"- {label}")
    else:
        print("- None")


def print_claim_mapping(path: str | Path) -> None:
    document = Document(path)
    mapping = map_claim_template(document)
    print("Claim template mapping:")
    print(f"- Claim table index: {mapping['claim_table_index']}")
    print(f"- Claim header row: {mapping['claim_header_row']}")
    print(f"- Claim row range: {mapping['claim_first_data_row']} to {mapping['claim_last_data_row']}")
    print("- Paragraph fields:")
    for field, index in mapping["paragraphs"].items():
        print(f"  - {field}: paragraph {index}")
    print("- Claim columns:")
    for field, index in mapping["claim_columns"].items():
        print(f"  - {field}: column {index}")


def print_register_mapping(path: str | Path) -> None:
    document = Document(path)
    mapping = map_register_template(document)
    print("Attendance register template mapping:")
    print(f"- Main table index: {mapping['table_index']}")
    print(f"- DATE row: {mapping['date_row']}")
    print(f"- TIME row: {mapping['time_row']}")
    print(f"- Student row range: {mapping['student_start_row']} to {mapping['student_end_row']}")
    print(f"- Surname column: {mapping['student_surname_column']}")
    print(f"- Initials column: {mapping['student_initials_column']}")
    print(f"- Student number column: {mapping['student_number_column']}")
    print(f"- Session columns: {mapping['session_columns']}")
    print(f"- Number of session columns: {len(mapping['session_columns'])}")
    print(f"- Footer paragraph range: {mapping['footer_paragraph_range'][0]} to {mapping['footer_paragraph_range'][1]}")
    print("- Paragraph fields:")
    for field, index in mapping["paragraphs"].items():
        print(f"  - {field}: paragraph {index}")


def _paragraph_diagnostics(paragraph) -> str:
    fmt = paragraph.paragraph_format
    tabs = len(fmt.tab_stops)
    return (
        f"alignment={paragraph.alignment}, "
        f"left_indent={fmt.left_indent}, "
        f"first_line_indent={fmt.first_line_indent}, "
        f"tab_stops={tabs}"
    )


def _merged_state(cell) -> str:
    tc_pr = cell._tc.tcPr
    if tc_pr is None:
        return "none"
    grid_span = tc_pr.find(qn("w:gridSpan"))
    v_merge = tc_pr.find(qn("w:vMerge"))
    parts = []
    if grid_span is not None:
        parts.append(f"gridSpan={grid_span.get(qn('w:val'))}")
    if v_merge is not None:
        parts.append(f"vMerge={v_merge.get(qn('w:val')) or 'continue'}")
    return ", ".join(parts) if parts else "none"


def print_register_cell_mapping(path: str | Path) -> None:
    document = Document(path)
    mapping = map_register_template(document)
    table = document.tables[mapping["table_index"]]
    print_register_mapping(path)
    print("Attendance register detailed cell map:")
    print(f"- Detected surname column: {mapping['student_surname_column']}")
    print(f"- Detected initials column: {mapping['student_initials_column']}")
    print(f"- Detected student number column: {mapping['student_number_column']}")
    print(f"- Detected DATE row: {mapping['date_row']}")
    print(f"- Detected TIME row: {mapping['time_row']}")
    print(f"- Detected student row range: {mapping['student_start_row']} to {mapping['student_end_row']}")
    print(f"- Detected footer paragraph range: {mapping['footer_paragraph_range'][0]} to {mapping['footer_paragraph_range'][1]}")
    for row_index, row in enumerate(table.rows):
        for col_index, cell in enumerate(row.cells):
            text = cell.text.replace("\n", " / ").strip()
            if text or row_index <= mapping["time_row"] or row_index in {mapping["student_start_row"], mapping["student_end_row"]}:
                paragraph = cell.paragraphs[0]
                print(
                    f"  table 0 row {row_index} col {col_index}: "
                    f"text={text!r}; merged={_merged_state(cell)}; {_paragraph_diagnostics(paragraph)}"
                )


def print_claim_cell_mapping(path: str | Path) -> None:
    document = Document(path)
    mapping = map_claim_template(document)
    table = document.tables[mapping["claim_table_index"]]
    print_claim_mapping(path)
    print("Claim detailed cell/paragraph map:")
    print("- Likely claimant field paragraphs:")
    for field, index in mapping["paragraphs"].items():
        print(f"  - {field}: paragraph {index}: {document.paragraphs[index].text!r}")
    print("- Likely right-hand value paragraphs:")
    for field in ("qualification_budget", "personnel_tariff", "identity_paye", "address_tel"):
        index = mapping["paragraphs"][field]
        print(f"  - {field}: paragraph {index}: {document.paragraphs[index].text!r}")
    print("- Training-level option labels are drawing/textbox content in the template; paragraph label:")
    level_index = mapping["paragraphs"]["level"]
    print(f"  - level label paragraph {level_index}: {document.paragraphs[level_index].text!r}")
    print(
        f"- Claim row range: {mapping['claim_first_data_row']} to {mapping['claim_last_data_row']} "
        f"in table {mapping['claim_table_index']}"
    )
    print("- Claim table non-empty cells:")
    for row_index, row in enumerate(table.rows):
        for col_index, cell in enumerate(row.cells):
            text = cell.text.replace("\n", " / ").strip()
            if text:
                paragraph = cell.paragraphs[0]
                print(
                    f"  table {mapping['claim_table_index']} row {row_index} col {col_index}: "
                    f"text={text!r}; merged={_merged_state(cell)}; {_paragraph_diagnostics(paragraph)}"
                )


def main() -> None:
    parser = argparse.ArgumentParser(description="Inspect DOCX template structure.")
    parser.add_argument("--file", required=True)
    parser.add_argument("--map-claim", action="store_true")
    parser.add_argument("--map-register", action="store_true")
    parser.add_argument("--map-claim-cells", action="store_true")
    parser.add_argument("--map-register-cells", action="store_true")
    args = parser.parse_args()
    print_template_analysis(inspect_template(args.file))
    if args.map_claim:
        print_claim_mapping(args.file)
    if args.map_register:
        print_register_mapping(args.file)
    if args.map_claim_cells:
        print_claim_cell_mapping(args.file)
    if args.map_register_cells:
        print_register_cell_mapping(args.file)


if __name__ == "__main__":
    main()
