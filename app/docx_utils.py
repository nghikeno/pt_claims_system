from datetime import date, datetime
from decimal import Decimal
from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt


def create_document() -> Document:
    document = Document()
    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10)
    styles["Title"].font.name = "Arial"
    styles["Title"].font.size = Pt(18)
    for style_name in ("Heading 1", "Heading 2", "Heading 3"):
        styles[style_name].font.name = "Arial"
    return document


def apply_compact_document_style(document: Document, font_name: str = "Arial", font_size: int = 9) -> None:
    for style_name in ("Normal", "Title", "Heading 1", "Heading 2", "Heading 3"):
        try:
            style = document.styles[style_name]
        except KeyError:
            continue
        style.font.name = font_name
        style.font.size = Pt(font_size if style_name == "Normal" else max(font_size + 1, 10))
        if hasattr(style, "paragraph_format"):
            style.paragraph_format.space_before = Pt(0)
            style.paragraph_format.space_after = Pt(2)
            style.paragraph_format.line_spacing = 1


def set_a4_portrait(section) -> None:
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)


def set_a4_landscape(section) -> None:
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21)


def set_landscape(section) -> None:
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width


def set_narrow_margins(section) -> None:
    section.top_margin = Inches(0.35)
    section.bottom_margin = Inches(0.35)
    section.left_margin = Inches(0.35)
    section.right_margin = Inches(0.35)


def add_heading(document: Document, text: str, level: int = 1):
    return document.add_heading(text, level=level)


def add_paragraph_text(document: Document, text: str, bold: bool = False):
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = bold
    return paragraph


def add_compact_paragraph(container, text: str = "", bold: bool = False, align=None, size: int = 9):
    paragraph = container.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(1)
    paragraph.paragraph_format.line_spacing = 1
    if align is not None:
        paragraph.alignment = align
    run = paragraph.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(size)
    return paragraph


def add_warning(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(12)


def create_table(document: Document, headers: list[str], rows: list[list[str]]):
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        header_cells[index].text = header
        for paragraph in header_cells[index].paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(8)
        header_cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for row_values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row_values):
            cells[index].text = str(value)
            cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cells[index].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8)
    return table


def set_cell_text(cell, text: str, bold: bool = False, size: int = 8, align=None) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1
    if align is not None:
        paragraph.alignment = align
    run = paragraph.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(size)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def create_fixed_table(document: Document, rows: int, cols: int, style: str = "Table Grid"):
    table = document.add_table(rows=rows, cols=cols)
    table.style = style
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    return table


def repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def safe_filename_text(value: str) -> str:
    cleaned = re.sub(r"[^A-Za-z0-9]+", "-", str(value).upper()).strip("-")
    return cleaned or "ITEM"


def format_date(value) -> str:
    if isinstance(value, datetime):
        return value.strftime("%Y-%m-%d")
    if isinstance(value, date):
        return value.strftime("%Y-%m-%d")
    parsed = datetime.strptime(str(value), "%Y-%m-%d")
    return parsed.strftime("%Y-%m-%d")


def format_docx_date(value) -> str:
    if isinstance(value, datetime):
        parsed = value
    elif isinstance(value, date):
        parsed = datetime.combine(value, datetime.min.time())
    else:
        parsed = datetime.strptime(str(value), "%Y-%m-%d")
    return parsed.strftime("%d-%b-%y")


def format_month_year(year: int, month: int) -> str:
    return date(year, month, 1).strftime("%B %Y")


def format_time(value) -> str:
    text = str(value)
    return text[:5]


def format_time_range(start_time, end_time) -> str:
    return f"{format_time(start_time)}-{format_time(end_time)}"


def format_hours(value) -> str:
    return f"{float(value):.2f}"


def format_compact_hours(value) -> str:
    parsed = float(value)
    return str(int(parsed)) if parsed.is_integer() else f"{parsed:.2f}"


def format_whole_money(value) -> str:
    return str(int(round(float(value))))


def format_currency(value) -> str:
    return f"{Decimal(str(value)):.2f}"


def ensure_parent(path: str | Path) -> Path:
    output = Path(path)
    output.parent.mkdir(parents=True, exist_ok=True)
    return output
