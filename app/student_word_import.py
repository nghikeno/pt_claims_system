from __future__ import annotations

from dataclasses import dataclass, field
from pathlib import Path
import re
from typing import BinaryIO, Iterable

from docx import Document

from app.student_row_safety import is_time_range_text, suspicious_student_row_reason


BANK_PATTERNS = (
    "bank",
    "account number",
    "account holder",
    "branch code",
    "account type",
    "swift",
    "first national bank",
    "fnb",
)


@dataclass
class ParsedAttendanceSheet:
    source_name: str
    header: dict[str, str] = field(default_factory=dict)
    students: list[dict[str, str]] = field(default_factory=list)
    skipped_rows: list[dict[str, str]] = field(default_factory=list)
    warnings: list[str] = field(default_factory=list)


def clean_text(value: object) -> str:
    return re.sub(r"\s+", " ", "" if value is None else str(value)).strip()


def _normalise_label(value: str) -> str:
    return re.sub(r"[^a-z0-9]+", " ", value.casefold()).strip()


def _all_table_rows(document: Document) -> Iterable[list[str]]:
    for table in document.tables:
        for row in table.rows:
            yield [clean_text(cell.text) for cell in row.cells]


def _paragraph_text(document: Document) -> list[str]:
    return [clean_text(paragraph.text) for paragraph in document.paragraphs if clean_text(paragraph.text)]


def contains_bank_detail_text(text: str) -> bool:
    lowered = clean_text(text).casefold()
    return any(pattern in lowered for pattern in BANK_PATTERNS)


def _extract_header_from_cells(cells: list[str], header: dict[str, str]) -> None:
    label_map = {
        "faculty": "faculty",
        "department": "department",
        "course name": "course_name",
        "course code": "course_code",
        "group": "group_label",
        "name of lecturer": "lecturer_name",
        "staff nr": "lecturer_staff_number",
        "staff no": "lecturer_staff_number",
        "staff number": "lecturer_staff_number",
    }
    for index, cell in enumerate(cells):
        if not cell:
            continue
        if ":" in cell:
            label, value = cell.split(":", 1)
            normalized = _normalise_label(label)
            for prefix, key in label_map.items():
                if normalized.startswith(prefix) and clean_text(value):
                    header.setdefault(key, clean_text(value))
        normalized_cell = _normalise_label(cell)
        for label, key in label_map.items():
            if normalized_cell == label and index + 1 < len(cells) and clean_text(cells[index + 1]):
                header.setdefault(key, clean_text(cells[index + 1]))


def _student_number_from_cells(cells: list[str]) -> tuple[int | None, str]:
    candidates: list[tuple[int, str]] = []
    row_text = " ".join(cells)
    for index, cell in enumerate(cells):
        if re.fullmatch(r"\d{1,2}[-/]\d{1,2}[-/]\d{2,4}", clean_text(cell)):
            continue
        if is_time_range_text(cell):
            continue
        compact = re.sub(r"\D", "", cell)
        if len(compact) >= 7 and suspicious_student_row_reason(compact, row_text=row_text) is None:
            candidates.append((index, compact))
    if candidates:
        return candidates[0]
    return None, ""


def _is_row_number(value: str) -> bool:
    return bool(re.fullmatch(r"\d{1,3}\.?", clean_text(value)))


def _clean_initials(value: str) -> str:
    return clean_text(value).replace(" ", "").replace(".", "")


def _split_name_cell(value: str) -> tuple[str, str]:
    value = clean_text(value)
    if "," in value:
        surname, initials = value.split(",", 1)
        return clean_text(surname), clean_text(initials).replace(" ", "")
    parts = value.split()
    if len(parts) >= 2 and re.fullmatch(r"[A-Za-z.]{1,8}", parts[-1]):
        return " ".join(parts[:-1]), parts[-1].replace(".", "")
    return value, ""


def _looks_like_student_row(cells: list[str]) -> bool:
    joined = " ".join(cells).casefold()
    reason = suspicious_student_row_reason(row_text=" ".join(cells))
    if reason is not None:
        return False
    _index, student_number = _student_number_from_cells(cells)
    return bool(student_number)


def _extract_student_from_row(cells: list[str]) -> dict[str, str] | None:
    if not _looks_like_student_row(cells):
        return None
    number_index, student_number = _student_number_from_cells(cells)
    if number_index is None:
        return None

    before_number = [clean_text(cell) for cell in cells[:number_index] if clean_text(cell)]
    if before_number and _is_row_number(before_number[0]):
        before_number = before_number[1:]

    surname = ""
    initials = ""
    if len(before_number) >= 2:
        surname = clean_text(before_number[-2])
        initials = _clean_initials(before_number[-1])
    elif len(before_number) == 1:
        surname, initials = _split_name_cell(before_number[0])

    if surname and not _is_row_number(surname):
        student = {
            "student_number": student_number,
            "surname": surname,
            "initials": initials,
            "full_name": " ".join(part for part in (surname, initials) if part),
        }
        if suspicious_student_row_reason(
            student["student_number"],
            student["surname"],
            student["initials"],
            student["full_name"],
            " ".join(cells),
        ):
            return None
        return student

    candidates = [cell for index, cell in enumerate(cells) if index != number_index and cell]
    if candidates and _is_row_number(candidates[0]):
        candidates = candidates[1:]
    if not candidates:
        return None

    surname = ""
    initials = ""
    if len(candidates) >= 2 and len(candidates[1]) <= 10 and re.fullmatch(r"[A-Za-z. ]+", candidates[1]):
        surname = clean_text(candidates[0])
        initials = clean_text(candidates[1]).replace(" ", "").replace(".", "")
    else:
        surname, initials = _split_name_cell(candidates[0])

    if not surname:
        return None
    student = {
        "student_number": student_number,
        "surname": surname,
        "initials": initials,
        "full_name": " ".join(part for part in (surname, initials) if part),
    }
    if suspicious_student_row_reason(
        student["student_number"],
        student["surname"],
        student["initials"],
        student["full_name"],
        " ".join(cells),
    ):
        return None
    return student


def parse_attendance_docx(path_or_file: str | Path | BinaryIO, source_name: str | None = None) -> ParsedAttendanceSheet:
    document = Document(path_or_file)
    source = source_name or getattr(path_or_file, "name", None) or str(path_or_file)
    parsed = ParsedAttendanceSheet(source_name=Path(str(source)).name)

    all_text = "\n".join(_paragraph_text(document))
    if contains_bank_detail_text(all_text):
        parsed.warnings.append("Bank details detected and ignored.")

    for paragraph in _paragraph_text(document):
        _extract_header_from_cells(re.split(r"\t+|\s{2,}", paragraph), parsed.header)

    seen_numbers: set[str] = set()
    for cells in _all_table_rows(document):
        row_text = " ".join(cells)
        if contains_bank_detail_text(row_text):
            parsed.warnings.append("Bank details detected and ignored.")
            continue
        _extract_header_from_cells(cells, parsed.header)
        suspicious_reason = suspicious_student_row_reason(row_text=row_text)
        if suspicious_reason:
            if any(cell for cell in cells):
                parsed.skipped_rows.append({"row_text": row_text, "reason": suspicious_reason})
            continue
        student = _extract_student_from_row(cells)
        if student is None:
            if any(cell for cell in cells) and not any(
                marker in row_text.casefold()
                for marker in ("nr", "student surname", "std nr", "signature", "name of lecturer", "date")
            ):
                parsed.skipped_rows.append({"row_text": row_text, "reason": "No valid student number or surname found."})
            continue
        if student["student_number"] in seen_numbers:
            parsed.warnings.append(f"Duplicate student number in uploaded file: {student['student_number']}")
        seen_numbers.add(student["student_number"])
        parsed.students.append(student)

    return parsed
