from copy import deepcopy
from datetime import date, datetime
from pathlib import Path
import shutil

import pandas as pd
from docx import Document
from docx.shared import Pt

from app.config import DOCX_TEMPLATES_DIR
from app.docx_utils import ensure_parent, format_compact_hours, format_time_range, format_whole_money


CLAIM_TEMPLATE_PATH = DOCX_TEMPLATES_DIR / "golden_claim_template.docx"
FORBIDDEN_SAMPLE_VALUES = [
    "Alvina",
    "Hailonga",
    "1009470",
    "97012300132",
    "14819389",
    "P.O Box 6951",
    "Ausspanplatz",
    "0816424142",
]


class MissingDocxTemplateError(FileNotFoundError):
    pass


class TemplateMappingError(ValueError):
    pass


def require_template(path: str | Path) -> Path:
    template_path = Path(path)
    if not template_path.exists():
        display = f"data/docx_templates/{template_path.name}"
        if template_path.name.startswith("golden_"):
            raise MissingDocxTemplateError(
                f"Missing golden DOCX template: {display}\n"
                "Copy the user-corrected template into data/docx_templates/ before generating documents."
            )
        raise MissingDocxTemplateError(
            f"Missing DOCX template: {display}\nPlease copy the official template into this folder."
        )
    return template_path


def _set_paragraph_text(paragraph, text: str, font_size: int | None = None) -> None:
    if not paragraph.runs:
        paragraph.add_run(str(text))
        return
    paragraph.runs[0].text = str(text)
    for run in paragraph.runs[1:]:
        run.text = ""


def _set_cell_text_preserve(cell, text: str, font_size: int | None = 9) -> None:
    paragraph = cell.paragraphs[0]
    _set_paragraph_text(paragraph, str(text), font_size=font_size)
    for extra in cell.paragraphs[1:]:
        _set_paragraph_text(extra, "")


def _paragraph_index(document: Document, needle: str) -> int:
    needle_lower = needle.lower()
    for index, paragraph in enumerate(document.paragraphs):
        if needle_lower in paragraph.text.lower():
            return index
    raise TemplateMappingError(f"Could not find paragraph containing: {needle}")


def _format_claim_date(value) -> str:
    if isinstance(value, datetime):
        parsed = value
    elif isinstance(value, date):
        parsed = datetime.combine(value, datetime.min.time())
    else:
        parsed = datetime.strptime(str(value), "%Y-%m-%d")
    return parsed.strftime("%d/%m/%Y")


def _course_post(sessions_df: pd.DataFrame) -> str:
    return " & ".join(sorted(set(sessions_df["course_code"].astype(str))))


def _faculty_department(sessions_df: pd.DataFrame) -> str:
    departments = sorted(set(sessions_df["department"].astype(str)))
    return " / ".join(departments)


def map_claim_template(document: Document) -> dict:
    if not document.tables:
        raise TemplateMappingError("Claim template does not contain a claim table")
    table_index, table = max(
        enumerate(document.tables),
        key=lambda item: len(item[1].rows) * max((len(row.cells) for row in item[1].rows), default=0),
    )
    header_row = None
    for index, row in enumerate(table.rows):
        row_text = " ".join(cell.text for cell in row.cells).lower()
        if "lecture" in row_text and "office use" in row_text:
            header_row = index
            break
    if header_row is None:
        raise TemplateMappingError("Could not identify claim table header row")
    return {
        "paragraphs": {
            "claimant": _paragraph_index(document, "Name & Surname"),
            "qualification_budget": _paragraph_index(document, "Highest qualification"),
            "personnel_tariff": _paragraph_index(document, "Personnel Number"),
            "identity_paye": _paragraph_index(document, "Identity / Passport Number"),
            "address_tel": _paragraph_index(document, "Address"),
            "level": _paragraph_index(document, "Level of training offered"),
            "hours": _paragraph_index(document, "HR"),
            "course_faculty": _paragraph_index(document, "Course/Post"),
        },
        "claim_table_index": table_index,
        "claim_header_row": header_row,
        "claim_first_data_row": header_row + 1,
        "claim_last_data_row": len(table.rows) - 1,
        "claim_columns": {
            "number": 0,
            "date": 1,
            "lecture": 2,
            "consultation": 3,
            "department_board": 4,
            "time": 5,
            "hours": 6,
            "tariff": 7,
            "cents": 8,
            "office": 9,
        },
    }


def _append_row_like(table, template_row_index: int):
    new_tr = deepcopy(table.rows[template_row_index]._tr)
    table._tbl.append(new_tr)
    return table.rows[-1]


def _delete_row(table, row_index: int) -> None:
    row = table.rows[row_index]
    table._tbl.remove(row._tr)


def _clear_claim_row(row) -> None:
    for cell in row.cells:
        _set_cell_text_preserve(cell, "")


def _fill_claim_particulars(document: Document, sessions_df: pd.DataFrame, mapping: dict) -> None:
    first = sessions_df.iloc[0]
    total_hours = round(float(sessions_df["hours"].sum()), 2)
    budget_allocations = ", ".join(sorted(set(sessions_df["budget_allocation"].astype(str))))
    p = document.paragraphs
    idx = mapping["paragraphs"]

    _set_paragraph_text(p[idx["claimant"]], f"Name & Surname: {first['lecturer_name']}", font_size=10)
    _set_paragraph_text(
        p[idx["qualification_budget"]],
        f"\tHighest qualification: {first['highest_qualification']}\t \tBudget Allocation: {budget_allocations}",
        font_size=10,
    )
    _set_paragraph_text(
        p[idx["personnel_tariff"]],
        f"\tPersonnel Number: {first['staff_number']}\t\t\t\t\tTariff per hour: {format_whole_money(first['tariff_per_hour'])}",
        font_size=10,
    )
    _set_paragraph_text(
        p[idx["identity_paye"]],
        f"\tIdentity / Passport Number: {first['id_or_passport_number']}\t\t\tPAYE No.: {first['paye_number']}",
        font_size=10,
    )
    _set_paragraph_text(
        p[idx["address_tel"]],
        f"Address: {first['physical_address']}\t\tTel. no.:\t {first['contact_number']}",
        font_size=10,
    )
    _set_paragraph_text(p[idx["level"]], "Level of training offered (indicate with an X):", font_size=10)
    _set_paragraph_text(
        p[idx["hours"]],
        f"PARTICULARS OF CLAIM\t\t\t\t\tHR – TOTAL HOURS CLAIMED: {format_compact_hours(total_hours)}",
    )
    _set_paragraph_text(
        p[idx["course_faculty"]],
        f"\nCourse/Post: {_course_post(sessions_df)}\t   Faculty/Department: {_faculty_department(sessions_df)}",
    )


def _fill_claim_rows(document: Document, sessions_df: pd.DataFrame, mapping: dict) -> None:
    table = document.tables[mapping["claim_table_index"]]
    first_row = mapping["claim_first_data_row"]
    last_row = mapping["claim_last_data_row"]
    cols = mapping["claim_columns"]

    sorted_df = sessions_df.sort_values(["group_name", "session_date", "start_time"])
    records: list[dict] = []
    for group_name, group_df in sorted_df.groupby("group_name", sort=True):
        for number, row in enumerate(group_df.to_dict("records"), start=1):
            records.append(
                {
                    "number": str(number),
                    "date": _format_claim_date(row["session_date"]),
                    "lecture": "X",
                    "consultation": group_name if number == 1 else "",
                    "department_board": "",
                    "time": format_time_range(row["start_time"], row["end_time"]),
                    "hours": format_compact_hours(row["hours"]),
                    "tariff": format_whole_money(row["tariff_per_hour"]),
                    "cents": "00",
                    "office": "",
                }
            )

    existing_capacity = last_row - first_row + 1
    target_rows = len(records) + 2
    if target_rows > existing_capacity:
        source_row = first_row
        for _ in range(target_rows - existing_capacity):
            _append_row_like(table, source_row)
        last_row = len(table.rows) - 1
    elif target_rows < existing_capacity:
        first_delete = first_row + target_rows
        for row_index in range(last_row, first_delete - 1, -1):
            _delete_row(table, row_index)
        last_row = first_row + target_rows - 1

    for row_index in range(first_row, last_row + 1):
        _clear_claim_row(table.rows[row_index])

    for offset, record in enumerate(records):
        row = table.rows[first_row + offset]
        for key, column_index in cols.items():
            _set_cell_text_preserve(row.cells[column_index], record[key])


def generate_template_claim_form(
    sessions_df: pd.DataFrame,
    output_path: str | Path,
    year: int,
    month: int,
    warning: bool = False,
    template_path: str | Path = CLAIM_TEMPLATE_PATH,
    strict: bool = True,
) -> Path:
    if sessions_df.empty:
        raise ValueError("Cannot generate claim form without sessions")

    template = require_template(template_path)
    output = ensure_parent(output_path)
    shutil.copy2(template, output)
    document = Document(output)
    mapping = map_claim_template(document)
    _fill_claim_particulars(document, sessions_df, mapping)
    _fill_claim_rows(document, sessions_df, mapping)
    document.save(output)
    return output
