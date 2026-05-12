from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document


REPORT_PATH = Path("docs") / "template_layout_diagnostics_claim.md"
RISKY_PLACEHOLDERS = (
    "{{ highest_qualification }}",
    "{{ physical_address }}",
    "{{ id_or_passport_number }}",
    "{{ paye_number }}",
    "{{ contact_number }}",
)
LEFT_FIELDS = (
    "{{ lecturer_name }}",
    "{{ highest_qualification }}",
    "{{ staff_number }}",
    "{{ id_or_passport_number }}",
    "{{ physical_address }}",
)
RIGHT_FIELDS = (
    "{{ budget_allocation }}",
    "{{ tariff_per_hour }}",
    "{{ paye_number }}",
    "{{ contact_number }}",
)


def _paragraph_table_cell_count(document: Document, text: str) -> int:
    count = 0
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if text in cell.text:
                    count += 1
    return count


def diagnose_template_layout(template_path: Path, report_path: Path = REPORT_PATH) -> dict:
    document = Document(template_path)
    risky_paragraphs = []
    multi_placeholder_paragraphs = []
    left_right_paragraphs = []

    for index, paragraph in enumerate(document.paragraphs):
        text = paragraph.text
        placeholders = [part for part in text.split() if part.startswith("{{") or part.endswith("}}")]
        placeholder_count = text.count("{{")
        has_left = any(field in text for field in LEFT_FIELDS)
        has_right = any(field in text for field in RIGHT_FIELDS)
        risky = [field for field in RISKY_PLACEHOLDERS if field in text]
        if placeholder_count > 1:
            multi_placeholder_paragraphs.append((index, text))
        if has_left and has_right:
            left_right_paragraphs.append((index, text))
        if risky:
            risky_paragraphs.append((index, risky, text))

    risky_in_tables = {
        placeholder: _paragraph_table_cell_count(document, placeholder)
        for placeholder in RISKY_PLACEHOLDERS
    }
    claimant_details_in_paragraphs = bool(risky_paragraphs)

    report_path.parent.mkdir(parents=True, exist_ok=True)
    lines = [
        "# Claim Template Layout Diagnostics",
        "",
        f"Template inspected: `{template_path}`",
        "",
        "## Summary",
        "",
        f"- Paragraphs: {len(document.paragraphs)}",
        f"- Tables: {len(document.tables)}",
        f"- Claimant details appear to be in ordinary paragraphs: {'Yes' if claimant_details_in_paragraphs else 'No'}",
        "",
        "## Paragraphs Containing More Than One Placeholder",
        "",
    ]
    if multi_placeholder_paragraphs:
        for index, text in multi_placeholder_paragraphs:
            lines.append(f"- Paragraph {index}: `{text}`")
    else:
        lines.append("- None detected.")

    lines.extend(["", "## Paragraphs Mixing Left And Right Claimant Fields", ""])
    if left_right_paragraphs:
        for index, text in left_right_paragraphs:
            lines.append(f"- Paragraph {index}: `{text}`")
    else:
        lines.append("- None detected.")

    lines.extend(["", "## Long-Replacement Reflow Risks", ""])
    if risky_paragraphs:
        for index, fields, text in risky_paragraphs:
            lines.append(f"- Paragraph {index}: {', '.join(fields)}")
            lines.append(f"  - Text: `{text}`")
    else:
        lines.append("- None detected.")

    lines.extend(["", "## Risky Placeholder Table-Cell Occurrences", ""])
    for placeholder, count in risky_in_tables.items():
        lines.append(f"- `{placeholder}`: {count}")

    lines.extend(
        [
            "",
            "## Assessment",
            "",
            "The detected claimant-detail placeholders are risky when they sit in ordinary paragraphs with tabs or spaces. "
            "Long replacement values can cause Word to reflow the paragraph, which may move right-hand fields such as "
            "Budget Allocation, Tariff per hour, PAYE No., and Tel. no. out of alignment.",
            "",
        ]
    )
    report_path.write_text("\n".join(lines), encoding="utf-8")
    return {
        "template_path": template_path,
        "report_path": report_path,
        "multi_placeholder_paragraphs": multi_placeholder_paragraphs,
        "left_right_paragraphs": left_right_paragraphs,
        "risky_paragraphs": risky_paragraphs,
        "claimant_details_in_paragraphs": claimant_details_in_paragraphs,
    }


def main() -> None:
    parser = argparse.ArgumentParser(description="Diagnose docxtpl claim template layout reflow risks.")
    parser.add_argument("--file", type=Path, required=True)
    args = parser.parse_args()
    result = diagnose_template_layout(args.file)
    print(f"Template inspected: {result['template_path']}")
    print(f"Diagnostic report written to: {result['report_path']}")
    print(f"Risky claimant-detail paragraphs: {len(result['risky_paragraphs'])}")
    print(f"Mixed left/right claimant paragraphs: {len(result['left_right_paragraphs'])}")


if __name__ == "__main__":
    main()
