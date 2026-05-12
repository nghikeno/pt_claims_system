from __future__ import annotations

import argparse
import shutil
from datetime import datetime
from pathlib import Path

from docx import Document
from docxtpl import DocxTemplate

from app_docxtpl.create_v2_templates import CLAIM_TEMPLATE_V2, REGISTER_TEMPLATE_V2, TEMPLATE_DIR, sha256


MANUAL_CLAIM_TEMPLATE_V2 = TEMPLATE_DIR / "manual_claim_template_v2.docx"
MANUAL_REGISTER_TEMPLATE_V2 = TEMPLATE_DIR / "manual_register_template_v2.docx"


class ManualTemplateError(RuntimeError):
    pass


def _dummy_context() -> dict:
    student = {
        "nr": "1",
        "surname": "DummySurname",
        "initials": "D.",
        "student_number": "900000000",
        "sig1": "",
        "sig2": "",
        "sig3": "",
        "sig4": "",
        "sig5": "",
    }
    claim_row = {
        "no": "1",
        "date": "01/01/2026",
        "activity": "X",
        "group_display": "Dummy Group",
        "meeting": "",
        "time_range": "08:00-09:00",
        "hours": "1",
        "rate": "1",
        "cents": "00",
        "office_use": "",
    }
    context = {
        "lecturer_title": "Ms",
        "lecturer_name": "Dummy Lecturer",
        "highest_qualification": "Dummy Qualification",
        "budget_allocation": "0000-0000",
        "staff_number": "000000",
        "tariff_per_hour": "1",
        "id_or_passport_number": "DUMMY-ID",
        "paye_number": "DUMMY-PAYE",
        "physical_address": "DUMMY ADDRESS",
        "contact_number": "0000000000",
        "level_part_time_mark": "X",
        "level_full_time_mark": "",
        "level_extra_curricular_mark": "",
        "total_hours": "1",
        "course_post": "DUMMY101",
        "faculty_department": "Dummy Department",
        "claim_rows": [claim_row],
        "faculty": "Dummy Faculty",
        "department": "Dummy Department",
        "course_name": "Dummy Course",
        "course_code": "DUMMY101",
        "group_name": "Dummy Group",
        "session_1_date": "01-01-26",
        "session_2_date": "",
        "session_3_date": "",
        "session_4_date": "",
        "session_5_date": "",
        "session_1_time": "08:00-09:00",
        "session_2_time": "",
        "session_3_time": "",
        "session_4_time": "",
        "session_5_time": "",
        "students": [student],
        "page_number": 1,
        "total_pages": 1,
    }
    return context


def require_manual_templates() -> None:
    for path in (MANUAL_CLAIM_TEMPLATE_V2, MANUAL_REGISTER_TEMPLATE_V2):
        if not path.exists():
            raise FileNotFoundError(
                f"Missing manual template: {path}\n"
                "Copy the user-edited manual template into data/docx_templates_v2/ before running Phase 3.2."
            )


def validate_docxtpl_template(path: Path) -> None:
    try:
        doc = DocxTemplate(path)
        doc.render(_dummy_context(), autoescape=True)
    except Exception as exc:
        raise ManualTemplateError(
            f"docxtpl validation failed for {path}: {exc}\n"
            "The likely cause is an invalid Jinja/docxtpl tag in the manual template. "
            "Open the named template and check table-row tags such as {%tr for ... %} and {%tr endfor %}."
        ) from exc


def _file_modified_time(path: Path) -> str:
    return datetime.fromtimestamp(path.stat().st_mtime).isoformat(sep=" ", timespec="seconds")


def _file_size(path: Path) -> int:
    return path.stat().st_size


def template_contains_text(path: Path, text: str) -> bool:
    document = Document(path)
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return text in "\n".join(parts)


def extract_docx_text(path: Path, limit: int | None = None) -> str:
    document = Document(path)
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    text = "\n".join(parts)
    return text[:limit] if limit is not None else text


def prepare_manual_templates_for_render(validate: bool = True, force: bool = True) -> dict:
    require_manual_templates()
    manual_claim_hash = sha256(MANUAL_CLAIM_TEMPLATE_V2)
    manual_register_hash = sha256(MANUAL_REGISTER_TEMPLATE_V2)
    shutil.copy2(MANUAL_CLAIM_TEMPLATE_V2, CLAIM_TEMPLATE_V2)
    shutil.copy2(MANUAL_REGISTER_TEMPLATE_V2, REGISTER_TEMPLATE_V2)
    render_claim_hash = sha256(CLAIM_TEMPLATE_V2)
    render_register_hash = sha256(REGISTER_TEMPLATE_V2)
    if manual_claim_hash != render_claim_hash:
        raise RuntimeError(
            "Manual claim template copy failed: render template hash does not match manual source hash."
        )
    if manual_register_hash != render_register_hash:
        raise RuntimeError(
            "Manual register template copy failed: render template hash does not match manual source hash."
        )
    if validate:
        validate_docxtpl_template(CLAIM_TEMPLATE_V2)
        validate_docxtpl_template(REGISTER_TEMPLATE_V2)

    manual_claim_hash_after = sha256(MANUAL_CLAIM_TEMPLATE_V2)
    manual_register_hash_after = sha256(MANUAL_REGISTER_TEMPLATE_V2)
    if manual_claim_hash != manual_claim_hash_after or manual_register_hash != manual_register_hash_after:
        raise RuntimeError("Manual template source hash changed while preparing render templates.")

    return {
        "manual_paths": {
            "claim": MANUAL_CLAIM_TEMPLATE_V2,
            "register": MANUAL_REGISTER_TEMPLATE_V2,
        },
        "render_paths": {
            "claim": CLAIM_TEMPLATE_V2,
            "register": REGISTER_TEMPLATE_V2,
        },
        "before": {
            "claim": manual_claim_hash,
            "register": manual_register_hash,
        },
        "after": {
            "claim": manual_claim_hash_after,
            "register": manual_register_hash_after,
        },
        "manual_claim_path": MANUAL_CLAIM_TEMPLATE_V2,
        "manual_claim_hash": manual_claim_hash,
        "render_claim_path": CLAIM_TEMPLATE_V2,
        "render_claim_hash": render_claim_hash,
        "manual_register_path": MANUAL_REGISTER_TEMPLATE_V2,
        "manual_register_hash": manual_register_hash,
        "render_register_path": REGISTER_TEMPLATE_V2,
        "render_register_hash": render_register_hash,
        "manual_claim_modified_time": _file_modified_time(MANUAL_CLAIM_TEMPLATE_V2),
        "manual_register_modified_time": _file_modified_time(MANUAL_REGISTER_TEMPLATE_V2),
        "render_claim_modified_time": _file_modified_time(CLAIM_TEMPLATE_V2),
        "render_register_modified_time": _file_modified_time(REGISTER_TEMPLATE_V2),
        "manual_claim_size": _file_size(MANUAL_CLAIM_TEMPLATE_V2),
        "manual_register_size": _file_size(MANUAL_REGISTER_TEMPLATE_V2),
        "render_claim_size": _file_size(CLAIM_TEMPLATE_V2),
        "render_register_size": _file_size(REGISTER_TEMPLATE_V2),
    }


def diagnose_manual_templates() -> dict:
    info = prepare_manual_templates_for_render(validate=False, force=True)
    print("Manual template diagnostic")
    for label in ("claim", "register"):
        manual_path = info[f"manual_{label}_path"]
        render_path = info[f"render_{label}_path"]
        manual_hash = info[f"manual_{label}_hash"]
        render_hash = info[f"render_{label}_hash"]
        print(f"{label.title()} manual path: {manual_path}")
        print(f"{label.title()} manual exists: {manual_path.exists()}")
        print(f"{label.title()} manual SHA256: {manual_hash}")
        print(f"{label.title()} manual modified time: {info[f'manual_{label}_modified_time']}")
        print(f"{label.title()} manual size: {info[f'manual_{label}_size']}")
        print(f"{label.title()} render path: {render_path}")
        print(f"{label.title()} render exists: {render_path.exists()}")
        print(f"{label.title()} render SHA256: {render_hash}")
        print(f"{label.title()} render modified time: {info[f'render_{label}_modified_time']}")
        print(f"{label.title()} render size: {info[f'render_{label}_size']}")
        print(f"{label.title()} hashes match after copy: {manual_hash == render_hash}")
    return info


def main() -> None:
    parser = argparse.ArgumentParser(description="Diagnose manual docxtpl template sources.")
    parser.add_argument("--diagnose", action="store_true", help="Print manual/render template hash diagnostics.")
    args = parser.parse_args()
    if args.diagnose:
        diagnose_manual_templates()
    else:
        parser.print_help()


if __name__ == "__main__":
    main()
