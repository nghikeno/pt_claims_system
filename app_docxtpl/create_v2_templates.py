from __future__ import annotations

import argparse
import hashlib
from pathlib import Path
import shutil

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt


TEMPLATE_DIR = Path("data") / "docx_templates_v2"
CLAIM_TEMPLATE_V2 = TEMPLATE_DIR / "claim_template_v2.docx"
REGISTER_TEMPLATE_V2 = TEMPLATE_DIR / "attendance_register_template_v2.docx"
USER_CLAIM_SOURCE = TEMPLATE_DIR / "user_claim_source.docx"
USER_REGISTER_SOURCE = TEMPLATE_DIR / "user_register_source.docx"


def _set_normal_style(document: Document) -> None:
    style = document.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(9)


def _cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = text
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = bold
            run.font.size = Pt(8)


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest().upper()


def _set_paragraph_text(paragraph, text: str) -> None:
    if not paragraph.runs:
        paragraph.add_run(text)
        return
    paragraph.runs[0].text = text
    for run in paragraph.runs[1:]:
        run.text = ""
    for extra in paragraph._p.xpath("./w:r[position()>1]"):
        pass


def _set_cell_text_preserve(cell, text: str) -> None:
    if not cell.paragraphs:
        cell.text = text
        return
    _set_paragraph_text(cell.paragraphs[0], text)
    for paragraph in cell.paragraphs[1:]:
        _set_paragraph_text(paragraph, "")


def _delete_row(table, row_index: int) -> None:
    table._tbl.remove(table.rows[row_index]._tr)


def _add_label_table(document: Document, rows: list[tuple[str, str]]) -> None:
    table = document.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for label, value in rows:
        cells = table.add_row().cells
        _cell_text(cells[0], label, bold=True)
        _cell_text(cells[1], value)


def create_claim_template(path: Path, overwrite: bool = False) -> Path:
    if path.exists() and not overwrite:
        return path
    path.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    _set_normal_style(document)
    section = document.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = Cm(1.2)

    title = document.add_paragraph("CLAIM FOR REMUNERATION BY PART-TIME LECTURERS / ACADEMIC STAFF MEMBERS")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].bold = True

    _add_label_table(
        document,
        [
            ("Title", "{{ lecturer_title }}"),
            ("Name & Surname", "{{ lecturer_name }}"),
            ("Highest qualification", "{{ highest_qualification }}"),
            ("Budget Allocation", "{{ budget_allocation }}"),
            ("Personnel Number", "{{ staff_number }}"),
            ("Tariff per hour", "{{ tariff_per_hour }}"),
            ("Identity / Passport Number", "{{ id_or_passport_number }}"),
            ("PAYE No.", "{{ paye_number }}"),
            ("Address", "{{ physical_address }}"),
            ("Tel. no.", "{{ contact_number }}"),
            (
                "Level of training offered",
                "Part-time [{{ level_part_time_mark }}]   "
                "Full-time [{{ level_full_time_mark }}]   "
                "Extra-curricular [{{ level_extra_curricular_mark }}]",
            ),
            ("HR - TOTAL HOURS CLAIMED", "{{ total_hours }}"),
            ("Course/Post", "{{ course_post }}"),
            ("Faculty/Department", "{{ faculty_department }}"),
        ],
    )

    document.add_paragraph("PARTICULARS OF CLAIM").runs[0].bold = True
    headers = [
        "No.",
        "Date",
        "Lecture/Tutorial/Admin.",
        "Consultation",
        "Dep./Board Meeting",
        "Time From-To",
        "Hours",
        "N$",
        "C",
        "FOR OFFICE USE",
    ]
    table = document.add_table(rows=4, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        _cell_text(table.rows[0].cells[index], header, bold=True)
    table.rows[1].cells[0].text = "{%tr for row in claim_rows %}"
    placeholders = [
        "{{ row.no }}",
        "{{ row.date }}",
        "{{ row.activity }}",
        "{{ row.group_display }}",
        "{{ row.meeting }}",
        "{{ row.time_range }}",
        "{{ row.hours }}",
        "{{ row.rate }}",
        "{{ row.cents }}",
        "{{ row.office_use }}",
    ]
    for index, placeholder in enumerate(placeholders):
        _cell_text(table.rows[2].cells[index], placeholder)
    table.rows[3].cells[0].text = "{%tr endfor %}"

    for text in [
        "Claimant's Signature: _______________________________    Date: _______________________",
        "Signature of Head of Department: _____________________    Date: _______________________",
        "Signature of Dean/Registrar: _________________________    Date: _______________________",
        "Processed by Payroll Department: _____________________    Date: _______________________",
        "VERY IMPORTANT: Claims older than 3 months will not be honoured.",
    ]:
        paragraph = document.add_paragraph(text)
        if text.startswith("VERY IMPORTANT"):
            paragraph.runs[0].bold = True
    document.save(path)
    return path


def create_register_template(path: Path, overwrite: bool = False) -> Path:
    if path.exists() and not overwrite:
        return path
    path.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    _set_normal_style(document)
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21)
    section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = Cm(1.0)

    heading = document.add_paragraph("CLASS ATTENDANCE SHEET")
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.runs[0].bold = True
    document.add_paragraph("FACULTY: {{ faculty }}")
    document.add_paragraph("DEPARTMENT: {{ department }}")
    document.add_paragraph("COURSE NAME: {{ course_name }}    COURSE CODE: {{ course_code }}    GROUP: {{ group_name }}")

    headers = ["NR.", "SURNAME", "INITIALS", "STD NR", "SIG 1", "SIG 2", "SIG 3", "SIG 4", "SIG 5"]
    table = document.add_table(rows=5, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        _cell_text(table.rows[0].cells[index], header, bold=True)
    date_cells = ["DATE:", "{{ session_1_date }}", "{{ session_2_date }}", "{{ session_3_date }}", "{{ session_4_date }}", "{{ session_5_date }}"]
    time_cells = ["TIME:", "{{ session_1_time }}", "{{ session_2_time }}", "{{ session_3_time }}", "{{ session_4_time }}", "{{ session_5_time }}"]
    _cell_text(table.rows[1].cells[3], date_cells[0], bold=True)
    for index, value in enumerate(date_cells[1:], start=4):
        _cell_text(table.rows[1].cells[index], value)
    _cell_text(table.rows[2].cells[3], time_cells[0], bold=True)
    for index, value in enumerate(time_cells[1:], start=4):
        _cell_text(table.rows[2].cells[index], value)
    table.rows[3].cells[0].text = "{%tr for student in students %}"
    student_placeholders = [
        "{{ student.nr }}",
        "{{ student.surname }}",
        "{{ student.initials }}",
        "{{ student.student_number }}",
        "{{ student.sig1 }}",
        "{{ student.sig2 }}",
        "{{ student.sig3 }}",
        "{{ student.sig4 }}",
        "{{ student.sig5 }}",
    ]
    for index, placeholder in enumerate(student_placeholders):
        _cell_text(table.rows[4].cells[index], placeholder)
    end_row = table.add_row()
    end_row.cells[0].text = "{%tr endfor %}"

    document.add_paragraph("")
    document.add_paragraph("{{ lecturer_name }}        __________________________        {{ staff_number }}        __________________")
    document.add_paragraph("NAME OF LECTURER          SIGNATURE                         STAFF NR.                 DATE")
    document.save(path)
    return path


def create_v2_templates(overwrite: bool = False) -> tuple[Path, Path]:
    print("WARNING: Simple v2 templates are for proof-of-concept only and are not HR-approved.")
    claim = create_claim_template(CLAIM_TEMPLATE_V2, overwrite=overwrite)
    register = create_register_template(REGISTER_TEMPLATE_V2, overwrite=overwrite)
    return claim, register


def _convert_claim_copy(path: Path) -> None:
    document = Document(path)
    for paragraph in document.paragraphs:
        text = paragraph.text
        if "Name & Surname" in text:
            _set_paragraph_text(paragraph, "Name & Surname: {{ lecturer_name }}")
        elif "Highest qualification" in text:
            _set_paragraph_text(paragraph, "\tHighest qualification: {{ highest_qualification }}\t \t              Budget Allocation: {{ budget_allocation }}")
        elif "Personnel Number" in text:
            _set_paragraph_text(paragraph, "\tPersonnel Number: {{ staff_number }}\t\t\t\t\tTariff per hour: {{ tariff_per_hour }}")
        elif "Identity / Passport Number" in text:
            _set_paragraph_text(paragraph, "\tIdentity / Passport Number: {{ id_or_passport_number }}\t\t\tPAYE No.: {{ paye_number }}")
        elif "Address" in text:
            _set_paragraph_text(paragraph, "Address: {{ physical_address }}\t\t                             Tel. no.: {{ contact_number }}")
        elif "Level of training offered" in text:
            _set_paragraph_text(paragraph, "Level of training offered (indicate with an X):               ")
        elif text.strip().startswith("PARTICULARS OF CLAIM") and "CLAIMANT" not in text:
            _set_paragraph_text(paragraph, "PARTICULARS OF CLAIM\t\t\t\t\tHR – TOTAL HOURS CLAIMED: {{ total_hours }}")
        elif "Course/Post" in text:
            _set_paragraph_text(paragraph, "\nCourse/Post: {{ course_post }}\t   Faculty/Department: {{ faculty_department }}")

    if not document.tables:
        raise ValueError("User claim source does not contain a claim table")
    table = document.tables[0]
    if len(table.rows) < 4:
        raise ValueError("User claim source claim table does not contain enough rows")
    _set_cell_text_preserve(table.rows[1].cells[0], "{%tr for row in claim_rows %}")
    for cell in table.rows[1].cells[1:]:
        _set_cell_text_preserve(cell, "")
    placeholders = [
        "{{ row.no }}",
        "{{ row.date }}",
        "{{ row.activity }}",
        "{{ row.group_display }}",
        "{{ row.meeting }}",
        "{{ row.time_range }}",
        "{{ row.hours }}",
        "{{ row.rate }}",
        "{{ row.cents }}",
        "{{ row.office_use }}",
    ]
    for index, placeholder in enumerate(placeholders):
        _set_cell_text_preserve(table.rows[2].cells[index], placeholder)
    _set_cell_text_preserve(table.rows[3].cells[0], "{%tr endfor %}")
    for cell in table.rows[3].cells[1:]:
        _set_cell_text_preserve(cell, "")
    for row_index in range(len(table.rows) - 1, 3, -1):
        _delete_row(table, row_index)
    document.save(path)


def _convert_register_copy(path: Path) -> None:
    document = Document(path)
    for paragraph in document.paragraphs:
        if paragraph.text.startswith("FACULTY:"):
            _set_paragraph_text(paragraph, "FACULTY:\t___{{ faculty }}______________________________________________________________________")
        elif paragraph.text.startswith("DEPARTMENT:"):
            _set_paragraph_text(paragraph, "DEPARTMENT: {{ department }}")
        elif paragraph.text.startswith("COURSE NAME:"):
            _set_paragraph_text(
                paragraph,
                "COURSE NAME: {{ course_name }}\t                     COURSE CODE: {{ course_code }}                                                                 GROUP:  {{ group_name }}",
            )
        elif "Demo Clean Lecturer" in paragraph.text or "200001" in paragraph.text:
            _set_paragraph_text(
                paragraph,
                "{{ lecturer_name }}\t          _______________________________\t\t                     {{ staff_number }}\t",
            )

    if not document.tables:
        raise ValueError("User register source does not contain an attendance table")
    table = document.tables[0]
    if len(table.rows) < 6:
        raise ValueError("User register source attendance table does not contain enough rows")
    date_values = ["DATE:", "{{ session_1_date }}", "{{ session_2_date }}", "{{ session_3_date }}", "{{ session_4_date }}", "{{ session_5_date }}"]
    time_values = ["TIME:", "{{ session_1_time }}", "{{ session_2_time }}", "{{ session_3_time }}", "{{ session_4_time }}", "{{ session_5_time }}"]
    _set_cell_text_preserve(table.rows[1].cells[3], date_values[0])
    _set_cell_text_preserve(table.rows[2].cells[3], time_values[0])
    for offset, value in enumerate(date_values[1:], start=4):
        _set_cell_text_preserve(table.rows[1].cells[offset], value)
    for offset, value in enumerate(time_values[1:], start=4):
        _set_cell_text_preserve(table.rows[2].cells[offset], value)

    _set_cell_text_preserve(table.rows[3].cells[0], "{%tr for student in students %}")
    for cell in table.rows[3].cells[1:]:
        _set_cell_text_preserve(cell, "")
    placeholders = [
        "{{ student.nr }}",
        "{{ student.surname }}",
        "{{ student.initials }}",
        "{{ student.student_number }}",
        "{{ student.sig1 }}",
        "{{ student.sig2 }}",
        "{{ student.sig3 }}",
        "{{ student.sig4 }}",
        "{{ student.sig5 }}",
    ]
    for index, placeholder in enumerate(placeholders):
        _set_cell_text_preserve(table.rows[4].cells[index], placeholder)
    _set_cell_text_preserve(table.rows[5].cells[0], "{%tr endfor %}")
    for cell in table.rows[5].cells[1:]:
        _set_cell_text_preserve(cell, "")
    for row_index in range(len(table.rows) - 1, 5, -1):
        _delete_row(table, row_index)
    document.save(path)


def create_v2_templates_from_user_sources(overwrite: bool = False) -> tuple[Path, Path, dict]:
    for source in (USER_CLAIM_SOURCE, USER_REGISTER_SOURCE):
        if not source.exists():
            raise FileNotFoundError(f"Missing user-approved source template: {source}")
    if not overwrite and (CLAIM_TEMPLATE_V2.exists() or REGISTER_TEMPLATE_V2.exists()):
        return CLAIM_TEMPLATE_V2, REGISTER_TEMPLATE_V2, {
            "before": {"claim": sha256(USER_CLAIM_SOURCE), "register": sha256(USER_REGISTER_SOURCE)},
            "after": {"claim": sha256(USER_CLAIM_SOURCE), "register": sha256(USER_REGISTER_SOURCE)},
        }

    before = {"claim": sha256(USER_CLAIM_SOURCE), "register": sha256(USER_REGISTER_SOURCE)}
    TEMPLATE_DIR.mkdir(parents=True, exist_ok=True)
    shutil.copy2(USER_CLAIM_SOURCE, CLAIM_TEMPLATE_V2)
    shutil.copy2(USER_REGISTER_SOURCE, REGISTER_TEMPLATE_V2)
    _convert_claim_copy(CLAIM_TEMPLATE_V2)
    _convert_register_copy(REGISTER_TEMPLATE_V2)
    after = {"claim": sha256(USER_CLAIM_SOURCE), "register": sha256(USER_REGISTER_SOURCE)}
    return CLAIM_TEMPLATE_V2, REGISTER_TEMPLATE_V2, {"before": before, "after": after}


def main() -> None:
    parser = argparse.ArgumentParser(description="Create experimental docxtpl v2 templates.")
    parser.add_argument("--overwrite", action="store_true")
    parser.add_argument("--from-user-sources", action="store_true")
    args = parser.parse_args()
    if args.from_user_sources:
        claim, register, hashes = create_v2_templates_from_user_sources(overwrite=args.overwrite)
        print(f"Claim source: {USER_CLAIM_SOURCE}")
        print(f"Claim template v2: {claim}")
        print(f"Register source: {USER_REGISTER_SOURCE}")
        print(f"Attendance register template v2: {register}")
        print("Source SHA256 before:")
        print(f"- Claim: {hashes['before']['claim']}")
        print(f"- Register: {hashes['before']['register']}")
        print("Source SHA256 after:")
        print(f"- Claim: {hashes['after']['claim']}")
        print(f"- Register: {hashes['after']['register']}")
    else:
        claim, register = create_v2_templates(overwrite=args.overwrite)
        print(f"Claim template v2: {claim}")
        print(f"Attendance register template v2: {register}")


if __name__ == "__main__":
    main()
