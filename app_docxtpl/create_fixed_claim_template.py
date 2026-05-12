from __future__ import annotations

import argparse
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


SOURCE_TEMPLATE = Path("data") / "docx_templates_v2" / "manual_claim_template_v2.docx"
FIXED_TEMPLATE = Path("data") / "docx_templates_v2" / "manual_claim_template_v2_fixed_table.docx"


CLAIMANT_ROWS = (
    ("Name & Surname: {{ lecturer_name }}", ""),
    ("Highest qualification: {{ highest_qualification }}", "Budget Allocation: {{ budget_allocation }}"),
    ("Personnel Number: {{ staff_number }}", "Tariff per hour: {{ tariff_per_hour }}"),
    ("Identity / Passport Number: {{ id_or_passport_number }}", "PAYE No.: {{ paye_number }}"),
    ("Address: {{ physical_address }}", "Tel. no.: {{ contact_number }}"),
)


def _delete_paragraph(paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def _insert_table_after(paragraph, table) -> None:
    paragraph._p.addnext(table._tbl)


def _remove_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "nil")


def _set_cell_width(cell, width_twips: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    width = tc_pr.first_child_found_in("w:tcW")
    if width is None:
        width = OxmlElement("w:tcW")
        tc_pr.append(width)
    width.set(qn("w:w"), str(width_twips))
    width.set(qn("w:type"), "dxa")


def _format_cell(cell, text: str, width_twips: int) -> None:
    cell.text = text
    _set_cell_width(cell, width_twips)
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        for run in paragraph.runs:
            run.font.size = Pt(10)


def _clone_paragraph_style(source, target) -> None:
    target.style = source.style
    target.paragraph_format.left_indent = source.paragraph_format.left_indent
    target.paragraph_format.right_indent = source.paragraph_format.right_indent
    target.paragraph_format.first_line_indent = source.paragraph_format.first_line_indent
    target.paragraph_format.space_before = source.paragraph_format.space_before
    target.paragraph_format.space_after = source.paragraph_format.space_after


def create_fixed_claim_template(
    source_path: Path = SOURCE_TEMPLATE,
    output_path: Path = FIXED_TEMPLATE,
) -> Path:
    if not source_path.exists():
        raise FileNotFoundError(f"Missing source manual claim template: {source_path}")

    document = Document(source_path)
    start_index = None
    end_index = None
    for index, paragraph in enumerate(document.paragraphs):
        if "PARTICULARS OF CLAIMANT" in paragraph.text:
            start_index = index
        if "Level of training offered" in paragraph.text and start_index is not None:
            end_index = index
            break
    if start_index is None or end_index is None:
        raise RuntimeError("Could not locate claimant-details paragraph range in the claim template.")

    anchor = document.paragraphs[start_index]
    template_paragraph = document.paragraphs[start_index + 1] if start_index + 1 < len(document.paragraphs) else anchor
    table = document.add_table(rows=len(CLAIMANT_ROWS), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    _remove_table_borders(table)
    left_width = int(Inches(4.8).twips)
    right_width = int(Inches(3.0).twips)
    for row, (left_text, right_text) in zip(table.rows, CLAIMANT_ROWS):
        _format_cell(row.cells[0], left_text, left_width)
        _format_cell(row.cells[1], right_text, right_width)
        _clone_paragraph_style(template_paragraph, row.cells[0].paragraphs[0])
        _clone_paragraph_style(template_paragraph, row.cells[1].paragraphs[0])

    # Preserve the source geometry of the table in its new location, then remove the appended original.
    table_xml = deepcopy(table._tbl)
    table._tbl.getparent().remove(table._tbl)
    anchor._p.addnext(table_xml)

    for paragraph in list(document.paragraphs[start_index + 1 : end_index]):
        _delete_paragraph(paragraph)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    return output_path


def main() -> None:
    parser = argparse.ArgumentParser(description="Create an experimental fixed-table v2 claim template.")
    parser.add_argument("--source", type=Path, default=SOURCE_TEMPLATE)
    parser.add_argument("--output", type=Path, default=FIXED_TEMPLATE)
    args = parser.parse_args()
    output = create_fixed_claim_template(args.source, args.output)
    print(f"Fixed-table claim template created: {output}")


if __name__ == "__main__":
    main()
