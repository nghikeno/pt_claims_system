from pathlib import Path
from zipfile import ZipFile

from docx import Document
import pytest

from app.create_maria_pilot_workbook import create_maria_pilot_workbook
from app.dev_reset import dev_reset
from app.document_generator import output_directory
from app.import_master_data import import_master_data
from app.session_generator import generate_monthly_sessions
from app.template_claim_generator import MissingDocxTemplateError
from app.template_register_generator import generate_template_attendance_register_pack


def _doc_text(path: Path) -> str:
    document = Document(path)
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


def _delete_if_exists(path: Path) -> None:
    if path.exists():
        try:
            path.unlink()
        except PermissionError:
            return False
    return True


def _page_break_count(path: Path) -> int:
    xml = ZipFile(path).read("word/document.xml").decode("utf-8")
    return xml.count('w:type="page"')


def _combined_doc_text(paths: list[Path]) -> str:
    return "\n".join(_doc_text(path) for path in paths)


@pytest.fixture
def maria_sessions():
    dev_reset()
    import_master_data(create_maria_pilot_workbook())
    return generate_monthly_sessions(1008977, 2026, 4)


def test_missing_attendance_template_gives_clear_error(maria_sessions, tmp_path):
    with pytest.raises(
        MissingDocxTemplateError,
        match="Missing golden DOCX template: data/docx_templates/golden_attendance_register_template.docx",
    ):
        generate_template_attendance_register_pack(
            maria_sessions,
            tmp_path / "register.docx",
            2026,
            4,
            template_path=tmp_path / "golden_attendance_register_template.docx",
        )


def test_maria_template_attendance_pack_contains_dynamic_values_and_students(maria_sessions):
    out_dir = output_directory(2026, 4, "1008977")
    output = out_dir / "attendance_registers_1008977_2026_04.docx"
    _delete_if_exists(output)
    paths = generate_template_attendance_register_pack(maria_sessions, output, 2026, 4)

    assert paths
    assert not output.exists()
    text = _combined_doc_text(paths)
    assert "CLASS ATTENDANCE SHEET" in text
    assert "Maria Matias" in text
    assert "1008977" in text
    assert "CUS411S" in text
    assert "ICT521S" in text
    assert "CUS HORTICULTURE" in text
    assert "ICT Distance" in text
    assert "09-04-26" in text
    assert "09-04-2026" not in text
    assert "10:30-12:30" in text
    assert "PilotSurname001" in text
    assert "Register ID" not in text
    assert "LECTURER:" not in text
    assert text.count("28-03-2026") == 0

    document = Document(paths[0])
    first_table = document.tables[0]
    first_student_row = first_table.rows[3].cells
    assert first_student_row[1].text == "PilotSurname013"
    assert first_student_row[2].text == "P013"
    assert first_student_row[3].text == "910000013"
    assert first_student_row[1].text != first_student_row[2].text
    row_numbers = {row.cells[0].text.strip(".") for row in first_table.rows}
    assert not {"30", "31", "32", "33", "34", "35", "36", "37"} & row_numbers
    assert len(first_table.rows) == 16


def test_demo_register_uses_compact_department_dates_and_student_numbers(tmp_path):
    dev_reset()
    from app.document_generator import generate_monthly_documents

    output = output_directory(2026, 2, "200001") / "attendance_registers_200001_2026_02.docx"
    if _delete_if_exists(output):
        result = generate_monthly_documents(200001, 2026, 2, allow_clashes=False)
        paths = [Path(path) for path in result["attendance_paths"]]
    else:
        output = tmp_path / "attendance_registers_200001_2026_02.docx"
        sessions = generate_monthly_sessions(200001, 2026, 2)
        paths = generate_template_attendance_register_pack(sessions, output, 2026, 2)
    assert paths
    output = paths[0]
    assert output.exists()
    document = Document(output)
    text = _combined_doc_text(paths)
    first_table = document.tables[0]

    assert "DEPARTMENT: Informatics" in text
    assert "Informatics and Journalism" not in "\n".join(paragraph.text for paragraph in document.paragraphs[:9])
    assert "02-02-26" in text
    assert "16-02-26" in text
    assert "23-02-26" in text
    assert "02-02-2026" not in text
    assert "08:00-09:00" in text
    assert "900000001" in text
    assert "900000012" in text
    assert any(row.cells[3].text == "900000001" for row in first_table.rows)
    first_student_row = first_table.rows[3].cells
    assert first_student_row[1].text == "Amunyela"
    assert first_student_row[2].text == "A."
    assert first_student_row[3].text == "900000001"
    assert "Amunyela Amunyela" not in text
    assert "Amunyela Amunyela A." not in text
    row_numbers = {row.cells[0].text.strip(".") for row in first_table.rows}
    assert not {"30", "31", "32", "33", "34", "35", "36", "37"} & row_numbers
    assert "Demo Clean Lecturer" in text
    assert "200001" in text
    assert "Register ID" not in text
    assert "Demo Group A" in text
    assert "Demo Group B" in text
    assert "Demo Group C" in text
    assert all(_page_break_count(path) == 0 for path in paths)


def test_template_register_uses_template_session_column_capacity(maria_sessions):
    out_dir = output_directory(2026, 4, "1008977")
    output = out_dir / "attendance_registers_1008977_2026_04.docx"
    paths = generate_template_attendance_register_pack(maria_sessions, output, 2026, 4)

    attendance_tables = []
    for path in paths:
        document = Document(path)
        attendance_tables.extend(
            table
            for table in document.tables
            if table.rows and table.rows[0].cells[0].text == "NR."
        )
    assert attendance_tables
    assert all(len(table.columns) <= 9 for table in attendance_tables)
    template_columns = len(Document("data/docx_templates/golden_attendance_register_template.docx").tables[0].columns)
    assert all(len(table.columns) == template_columns for table in attendance_tables)
