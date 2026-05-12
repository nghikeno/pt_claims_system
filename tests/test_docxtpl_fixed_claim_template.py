from __future__ import annotations

from pathlib import Path

from docx import Document

from app.create_maria_pilot_workbook import EXPECTED_AMOUNT, EXPECTED_HOURS, EXPECTED_SESSIONS, create_maria_pilot_workbook
from app.dev_reset import dev_reset
from app.import_master_data import import_master_data
from app.session_generator import generate_monthly_sessions
from app_docxtpl.create_fixed_claim_template import (
    FIXED_TEMPLATE,
    SOURCE_TEMPLATE,
    create_fixed_claim_template,
)
from app_docxtpl.render_claim_v2 import render_claim_v2
from app_docxtpl.template_layout_diagnostics import diagnose_template_layout


def _doc_text(path: Path) -> str:
    document = Document(path)
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


def test_template_layout_diagnostics_detects_risky_claimant_paragraphs(tmp_path):
    risky_template = tmp_path / "risky_claim_template.docx"
    document = Document()
    document.add_paragraph("PARTICULARS OF CLAIMANT:")
    document.add_paragraph(
        "Highest qualification: {{ highest_qualification }}        Budget Allocation: {{ budget_allocation }}"
    )
    document.add_paragraph(
        "Identity / Passport Number: {{ id_or_passport_number }}        PAYE No.: {{ paye_number }}"
    )
    document.add_paragraph("Address: {{ physical_address }}        Tel. no.: {{ contact_number }}")
    document.save(risky_template)
    report_path = tmp_path / "diagnostics.md"
    result = diagnose_template_layout(risky_template, report_path)

    assert result["claimant_details_in_paragraphs"]
    assert result["risky_paragraphs"]
    assert result["left_right_paragraphs"]
    assert report_path.exists()
    assert "{{ highest_qualification }}" in report_path.read_text(encoding="utf-8")


def test_create_fixed_claim_template_creates_table_with_required_placeholders():
    output = create_fixed_claim_template()
    document = Document(output)
    text = _doc_text(output)

    assert output == FIXED_TEMPLATE
    assert output.exists()
    assert len(document.tables) >= 2
    for placeholder in (
        "{{ lecturer_name }}",
        "{{ highest_qualification }}",
        "{{ budget_allocation }}",
        "{{ staff_number }}",
        "{{ tariff_per_hour }}",
        "{{ id_or_passport_number }}",
        "{{ paye_number }}",
        "{{ physical_address }}",
        "{{ contact_number }}",
    ):
        assert placeholder in text


def test_render_maria_claim_with_fixed_table_template_succeeds(monkeypatch, tmp_path):
    dev_reset()
    import_master_data(create_maria_pilot_workbook())
    create_fixed_claim_template()

    def _tmp_generated_v2_directory(year: int, month: int, staff_number: str) -> Path:
        return tmp_path / "generated_v2" / str(year) / f"{month:02d}" / str(staff_number)

    monkeypatch.setattr("app_docxtpl.render_claim_v2.generated_v2_directory", _tmp_generated_v2_directory)
    output = render_claim_v2(1008977, 2026, 4, template_path=FIXED_TEMPLATE)
    text = _doc_text(output)
    sessions_df = generate_monthly_sessions(1008977, 2026, 4)

    for expected in (
        "Maria Matias",
        "Master of Science in Information Technology",
        "Budget Allocation",
        "0183-0102",
        "Tariff per hour",
        "460",
        "PAYE No.",
        "PLACEHOLDER-PAYE-1008977",
        "Tel. no.",
        "0810000000",
    ):
        assert expected in text
    assert "{{" not in text
    assert "}}" not in text
    assert len(sessions_df) == EXPECTED_SESSIONS
    assert round(float(sessions_df["hours"].sum()), 2) == EXPECTED_HOURS
    assert round(float(sessions_df["amount"].sum()), 2) == EXPECTED_AMOUNT
