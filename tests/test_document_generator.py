import shutil
import hashlib
from pathlib import Path

from docx import Document
import pandas as pd
import pytest

from app.create_maria_pilot_workbook import (
    EXPECTED_AMOUNT,
    EXPECTED_HOURS,
    EXPECTED_SESSIONS,
    create_maria_pilot_workbook,
)
from app.dev_reset import dev_reset
from app.document_generator import generate_monthly_documents, output_directory
from app.import_master_data import import_master_data
from app.seed_data import seed_database
from app.session_generator import generate_monthly_sessions
from app.validators import detect_clashes


GOLDEN_CLAIM = Path("data/docx_templates/golden_claim_template.docx")
GOLDEN_REGISTER = Path("data/docx_templates/golden_attendance_register_template.docx")


@pytest.fixture(autouse=True)
def seeded_database(tmp_path, monkeypatch):
    monkeypatch.setattr("app.document_generator.GENERATED_DIR", tmp_path / "generated")
    seed_database()
    out_dir = output_directory(2026, 2, "100718")
    if out_dir.exists():
        shutil.rmtree(out_dir)


def _doc_text(path: Path) -> str:
    document = Document(path)
    paragraphs = [paragraph.text for paragraph in document.paragraphs]
    table_text = []
    for table in document.tables:
        for row in table.rows:
            table_text.extend(cell.text for cell in row.cells)
    return "\n".join(paragraphs + table_text)


def _attendance_text(result: dict) -> str:
    return "\n".join(_doc_text(Path(path)) for path in result["attendance_paths"])


def _sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def test_verification_checklist_excel_file_is_created_when_clashes_block_docx():
    result = generate_monthly_documents(1, 2026, 2, allow_clashes=False)

    assert result["documents_generated"] is False
    assert Path(result["verification_path"]).exists()


def test_if_clashes_exist_without_allow_clashes_docx_files_are_not_generated():
    result = generate_monthly_documents(1, 2026, 2, allow_clashes=False)

    assert not Path(result["attendance_path"]).exists()
    assert not Path(result["claim_path"]).exists()


def test_if_clashes_exist_with_allow_clashes_docx_files_are_generated_in_template_mode():
    result = generate_monthly_documents(1, 2026, 2, allow_clashes=True)

    claim_path = Path(result["claim_path"])
    assert result["attendance_paths"]
    assert all(Path(path).exists() for path in result["attendance_paths"])
    assert claim_path.exists()
    assert "Register ID" not in _attendance_text(result)


def test_no_clash_lecturer_generates_documents_without_allow_clashes():
    dev_reset()
    claim_hash_before = _sha256(GOLDEN_CLAIM)
    register_hash_before = _sha256(GOLDEN_REGISTER)

    result = generate_monthly_documents(200001, 2026, 2, allow_clashes=False)

    assert result["documents_generated"] is True
    assert result["attendance_paths"]
    assert all(Path(path).exists() for path in result["attendance_paths"])
    assert all(Path(path).name.startswith("register_200001_") for path in result["attendance_paths"])
    assert Path(result["claim_path"]).exists()
    assert Path(result["verification_path"]).exists()
    assert result["layout_mode"] == "template"
    assert _sha256(GOLDEN_CLAIM) == claim_hash_before
    assert _sha256(GOLDEN_REGISTER) == register_hash_before
    assert not Path(result["attendance_path"]).exists()
    assert all(Path(path).parent.name == "registers" for path in result["attendance_paths"])


def test_generated_layout_mode_still_works():
    dev_reset()
    result = generate_monthly_documents(200001, 2026, 2, allow_clashes=False, layout_mode="generated")

    assert result["documents_generated"] is True
    assert result["layout_mode"] == "generated"
    assert Path(result["attendance_path"]).exists()
    assert Path(result["claim_path"]).exists()


def test_clean_attendance_register_contains_required_content_and_students():
    dev_reset()
    result = generate_monthly_documents(200001, 2026, 2, allow_clashes=False)

    text = _attendance_text(result)

    assert "Monthly Attendance Register Pack" not in text
    assert "CLASS ATTENDANCE SHEET" in text
    assert "Demo Clean Lecturer" in text
    assert "200001" in text
    assert "CUS411S" in text
    assert "Demo Group A" in text
    assert "Demo Group B" in text
    assert "Demo Group C" in text
    assert "Amunyela" in text
    assert "Amutenya" in text
    assert "Andima" in text
    assert "NAME OF LECTURER" in text
    assert "SIGNATURE" in text
    assert "STAFF NR." in text
    assert "DATE" in text
    assert "Register ID" not in text
    assert "LECTURER:" not in text
    assert len(result["attendance_paths"]) == 3
    for path in result["attendance_paths"]:
        register_text = _doc_text(Path(path))
        groups_present = [group for group in ("Demo Group A", "Demo Group B", "Demo Group C") if group in register_text]
        assert len(groups_present) == 1


def test_clean_claim_form_contains_required_content():
    dev_reset()
    result = generate_monthly_documents(200001, 2026, 2, allow_clashes=False)

    text = _doc_text(Path(result["claim_path"]))

    assert "CLAIM FOR REMUNERATION BY PART-TIME LECTURERS / ACADEMIC STAFF MEMBERS" in text
    assert "Demo Clean Lecturer" in text
    assert "200001" in text
    assert "0183-0102" in text
    assert "410" in text
    assert "CUS411S" in text
    assert "Total" in text or "TOTAL" in text
    assert "9" in text
    assert "TOTAL AMOUNT" not in text
    assert "Claimant’s Signature" in text
    assert "Signature of Head of Department" in text
    assert "Signature of Dean/Registrar" in text
    assert "Processed by Payroll Department" in text
    assert "Claims older than 3 months will not be honoured" in text
    assert "Level of training offered (indicate with an X): Part-time" not in text
    assert "Level of training offered (indicate with an X): Part-time X" not in text
    assert "Part-ti" not in text


def test_clean_register_students_are_sorted_by_surname_initials_student_number():
    dev_reset()
    result = generate_monthly_documents(200001, 2026, 2, allow_clashes=False)
    text = _attendance_text(result)

    assert text.index("Amunyela") < text.index("Haingura") < text.index("Iileka")
    assert text.index("Amutenya") < text.index("Hamukoto") < text.index("Iipinge")
    assert text.index("Andima") < text.index("Haufiku") < text.index("Iita")


def test_verification_checklist_contains_improved_summary_fields():
    dev_reset()
    result = generate_monthly_documents(200001, 2026, 2, allow_clashes=False)
    summary_df = pd.read_excel(result["verification_path"], sheet_name="Summary")
    checklist_df = pd.read_excel(result["verification_path"], sheet_name="Document Checklist")

    metrics = set(summary_df["Metric"])
    assert "Generated by system" in metrics
    assert "Generated timestamp" in metrics
    assert "Campus" in metrics
    assert "Number of excluded dates" in metrics
    assert "Generation status" in metrics
    assert "evidence_or_page_reference" in checklist_df.columns


def test_maria_claim_form_uses_institutional_layout_and_expected_totals():
    dev_reset()
    import_master_data(create_maria_pilot_workbook())
    result = generate_monthly_documents(1008977, 2026, 4, allow_clashes=False)
    sessions_df = generate_monthly_sessions(1008977, 2026, 4)
    clashes_df = detect_clashes(sessions_df)

    text = _doc_text(Path(result["claim_path"]))

    assert len(sessions_df) == EXPECTED_SESSIONS
    assert round(float(sessions_df["hours"].sum()), 2) == EXPECTED_HOURS
    assert round(float(sessions_df["amount"].sum()), 2) == EXPECTED_AMOUNT
    assert clashes_df.empty
    assert "Field\nValue" not in text[:500]
    assert "CLAIM FOR REMUNERATION BY PART-TIME LECTURERS / ACADEMIC STAFF MEMBERS" in text
    assert "PARTICULARS OF CLAIMANT" in text
    assert "HR – TOTAL HOURS CLAIMED: 94" in text
    assert "Course/Post" in text
    assert "Faculty/Department" in text
    assert "CUS HORTICULTURE" in text
    assert "CUS GROUP A" in text
    assert "ICT GREY" in text
    assert "ICT Distance" in text
    assert "08/04/2026" in text
    assert "10:30-12:30" in text
    assert "Claimant’s Signature" in text
    assert "Signature of Head of Department" in text
    assert "Signature of Dean/Registrar" in text
    assert "Processed by Payroll Department" in text
    assert "Claims older than 3 months will not be honoured" in text


def test_maria_attendance_register_uses_institutional_layout():
    dev_reset()
    import_master_data(create_maria_pilot_workbook())
    result = generate_monthly_documents(1008977, 2026, 4, allow_clashes=False)

    text = _attendance_text(result)
    register_tables = []
    for path in result["attendance_paths"]:
        document = Document(path)
        register_tables.extend(
            table
            for table in document.tables
            if table.rows and table.rows[0].cells[0].text == "NR."
        )

    assert "Monthly Attendance Register Pack" not in text[:300]
    assert "CLASS ATTENDANCE SHEET" in text
    assert "FACULTY" in text
    assert "DEPARTMENT" in text
    assert "COURSE NAME" in text
    assert "COURSE CODE" in text
    assert "GROUP" in text
    assert "DATE:" in text
    assert "TIME:" in text
    assert "09-04-26" in text
    assert "09-04-2026" not in text
    assert "10:30-12:30" in text
    assert "NAME OF LECTURER" in text
    assert "SIGNATURE" in text
    assert "STAFF NR." in text
    assert "Register ID" not in text
    assert "LECTURER:" not in text
    assert register_tables
    assert len(register_tables[0].columns) <= 9
    assert len(register_tables[0].rows) < 32
