from pathlib import Path
from zipfile import ZipFile
import html
import re

from docx import Document
import pytest

from app.create_maria_pilot_workbook import create_maria_pilot_workbook
from app.dev_reset import dev_reset
from app.document_generator import output_directory
from app.import_master_data import import_master_data
from app.session_generator import generate_monthly_sessions
from app.template_claim_generator import (
    FORBIDDEN_SAMPLE_VALUES,
    MissingDocxTemplateError,
    generate_template_claim_form,
)
from app.validators import detect_clashes


def _doc_text(path: Path) -> str:
    document = Document(path)
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    xml = ZipFile(path).read("word/document.xml").decode("utf-8")
    parts.extend(html.unescape(match) for match in re.findall(r"<w:t[^>]*>(.*?)</w:t>", xml))
    return "\n".join(parts)


def _delete_if_exists(path: Path) -> None:
    if path.exists():
        try:
            path.unlink()
        except PermissionError:
            return False
    return True


@pytest.fixture
def maria_sessions():
    dev_reset()
    import_master_data(create_maria_pilot_workbook())
    sessions_df = generate_monthly_sessions(1008977, 2026, 4)
    assert len(sessions_df) == 69
    assert round(float(sessions_df["hours"].sum()), 2) == 94.00
    assert round(float(sessions_df["amount"].sum()), 2) == 43240.00
    assert detect_clashes(sessions_df).empty
    return sessions_df


def test_missing_claim_template_gives_clear_error(maria_sessions, tmp_path):
    with pytest.raises(MissingDocxTemplateError, match="Missing golden DOCX template: data/docx_templates/golden_claim_template.docx"):
        generate_template_claim_form(
            maria_sessions,
            tmp_path / "claim.docx",
            2026,
            4,
            template_path=tmp_path / "golden_claim_template.docx",
        )


def test_maria_template_claim_is_created_and_scrubs_sample_values(maria_sessions):
    out_dir = output_directory(2026, 4, "1008977")
    output = out_dir / "claim_form_1008977_2026_04.docx"
    generate_template_claim_form(maria_sessions, output, 2026, 4)

    assert output.exists()
    text = _doc_text(output)
    assert "Maria Matias" in text
    assert "1008977" in text
    assert "0183-0102" in text
    assert "HR – TOTAL HOURS CLAIMED: 94" in text
    assert "CUS HORTICULTURE" in text
    assert "CUS GROUP A" in text
    assert "ICT GREY" in text
    assert "ICT Distance" in text
    assert "08/04/2026" in text
    assert "10:30-12:30" in text
    assert "460" in text
    assert "TOTAL AMOUNT" not in text
    assert "Field\nValue" not in text
    assert "Part-time X" not in text
    assert "Level of training offered (indicate with an X): Part-time" not in text
    assert "Level of training offered (indicate with an X):" in text
    assert "Budget Allocation" in text
    assert "Tariff per hour" in text
    assert "PAYE No." in text
    assert "Tel. no." in text
    assert len(Document(output).tables) == len(Document("data/docx_templates/golden_claim_template.docx").tables)
    assert len(Document(output).tables[0].rows) == 72
    for sample_value in FORBIDDEN_SAMPLE_VALUES:
        assert sample_value not in text


def test_clean_demo_claim_regenerated_file_has_clean_training_level_and_labels(tmp_path):
    dev_reset()
    from app.document_generator import generate_monthly_documents

    output = output_directory(2026, 2, "200001") / "claim_form_200001_2026_02.docx"
    if _delete_if_exists(output):
        result = generate_monthly_documents(200001, 2026, 2, allow_clashes=False)
        assert Path(result["claim_path"]) == output
    else:
        output = tmp_path / "claim_form_200001_2026_02.docx"
        sessions = generate_monthly_sessions(200001, 2026, 2)
        generate_template_claim_form(sessions, output, 2026, 2)
    assert output.exists()

    text = _doc_text(output)
    assert "Level of training offered (indicate with an X): Part-time" not in text
    assert "Level of training offered (indicate with an X): Part-time X" not in text
    assert not re.search(r"\bPart-ti\b(?!me)", text)
    assert "Level of training offered (indicate with an X):" in text
    assert "Part-time" in text
    assert "Full-time" in text
    assert "Extra-curricular" in text
    assert "Budget Allocation" in text
    assert "Tariff per hour" in text
    assert "PAYE No." in text
    assert "Tel. no." in text
    assert "TOTAL AMOUNT" not in text
    for sample_value in FORBIDDEN_SAMPLE_VALUES:
        assert sample_value not in text
