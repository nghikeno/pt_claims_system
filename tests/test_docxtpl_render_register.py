from __future__ import annotations

from pathlib import Path

from docx import Document

from app.create_maria_pilot_workbook import create_maria_pilot_workbook
from app.dev_reset import dev_reset
from app.import_master_data import import_master_data
from app_docxtpl.create_v2_templates import REGISTER_TEMPLATE_V2
from app_docxtpl.manual_templates import prepare_manual_templates_for_render
from app_docxtpl.render_register_v2 import render_register_v2


def _doc_text(path: Path) -> str:
    document = Document(path)
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


def test_manual_register_template_contains_placeholders_and_preserves_institutional_text():
    prepare_manual_templates_for_render(validate=True)
    text = _doc_text(REGISTER_TEMPLATE_V2)

    for placeholder in (
        "{{ lecturer_name }}",
        "{{ staff_number }}",
        "{{ group_name }}",
        "{{ session_1_date }}",
        "{{ session_1_time }}",
    ):
        assert placeholder in text
    assert "{%tr for student in students %}" in text
    assert "{%tr endfor %}" in text
    for label in (
        "CLASS ATTENDANCE SHEET",
        "FACULTY",
        "DEPARTMENT",
        "COURSE NAME",
        "COURSE CODE",
        "GROUP",
        "NAME OF LECTURER",
        "SIGNATURE",
        "STAFF NR.",
        "DATE",
    ):
        assert label in text


def test_maria_register_v2_renders_separate_group_files():
    dev_reset()
    prepare_manual_templates_for_render(validate=True)
    import_master_data(create_maria_pilot_workbook())

    outputs = render_register_v2(1008977, 2026, 4)

    assert outputs
    assert all(path.exists() for path in outputs)
    assert all(path.parent == Path("data/generated_v2/2026/04/1008977/registers") for path in outputs)
    combined_text = "\n".join(_doc_text(path) for path in outputs)
    assert "CLASS ATTENDANCE SHEET" in combined_text
    assert "Maria Matias" in combined_text
    assert "1008977" in combined_text
    assert "09-04-26" in combined_text
    assert "10:30-12:30" in combined_text
    assert "PilotSurname001" in combined_text
    assert "{{" not in combined_text
    assert "}}" not in combined_text
    assert "Amunyela" not in combined_text

    group_names = ["CUS HORTICULTURE", "CUS GROUP A", "CUS GROUP B", "ICT GREY", "ICT Distance"]
    for path in outputs:
        text = _doc_text(path)
        groups_present = [group for group in group_names if group in text]
        assert len(groups_present) <= 1


def test_clean_demo_register_v2_renders():
    dev_reset()
    prepare_manual_templates_for_render(validate=True)

    outputs = render_register_v2(200001, 2026, 2)
    text = "\n".join(_doc_text(path) for path in outputs)

    assert len(outputs) == 3
    assert "Demo Clean Lecturer" in text
    assert "200001" in text
    assert "Demo Group A" in text
    assert "900000001" in text
