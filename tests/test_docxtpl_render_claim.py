from __future__ import annotations

from pathlib import Path

from docx import Document

from app.create_maria_pilot_workbook import create_maria_pilot_workbook
from app.dev_reset import dev_reset
from app.import_master_data import import_master_data
from app_docxtpl.create_v2_templates import (
    CLAIM_TEMPLATE_V2,
    REGISTER_TEMPLATE_V2,
    create_claim_template,
    create_register_template,
    sha256,
)
from app_docxtpl.manual_templates import (
    MANUAL_CLAIM_TEMPLATE_V2,
    MANUAL_REGISTER_TEMPLATE_V2,
    prepare_manual_templates_for_render,
)
from app_docxtpl.render_claim_v2 import render_claim_v2


def _doc_text(path: Path) -> str:
    document = Document(path)
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


def test_create_v2_templates_creates_templates_with_placeholders(tmp_path):
    claim = create_claim_template(tmp_path / "claim_template_v2.docx", overwrite=True)
    register = create_register_template(tmp_path / "attendance_register_template_v2.docx", overwrite=True)

    claim_text = _doc_text(claim)
    register_text = _doc_text(register)
    assert "{{ lecturer_name }}" in claim_text
    assert "{%tr for row in claim_rows %}" in claim_text
    assert "{{ session_1_date }}" in register_text
    assert "{%tr for student in students %}" in register_text


def test_prepare_manual_templates_preserves_source_hashes():
    assert MANUAL_CLAIM_TEMPLATE_V2.exists()
    assert MANUAL_REGISTER_TEMPLATE_V2.exists()
    claim_before = sha256(MANUAL_CLAIM_TEMPLATE_V2)
    register_before = sha256(MANUAL_REGISTER_TEMPLATE_V2)

    info = prepare_manual_templates_for_render(validate=True)

    assert info["render_paths"]["claim"] == CLAIM_TEMPLATE_V2
    assert info["render_paths"]["register"] == REGISTER_TEMPLATE_V2
    assert sha256(MANUAL_CLAIM_TEMPLATE_V2) == claim_before == info["before"]["claim"] == info["after"]["claim"]
    assert sha256(MANUAL_REGISTER_TEMPLATE_V2) == register_before == info["before"]["register"] == info["after"]["register"]


def test_manual_claim_template_contains_placeholders_and_preserves_institutional_text():
    prepare_manual_templates_for_render(validate=True)
    text = _doc_text(CLAIM_TEMPLATE_V2)

    for placeholder in ("{{ lecturer_name }}", "{{ staff_number }}", "{{ total_hours }}", "{{ course_post }}"):
        assert placeholder in text
    assert "{%tr for row in claim_rows %}" in text
    assert "{%tr endfor %}" in text
    assert "CLAIM FOR REMUNERATION BY PART-TIME LECTURERS / ACADEMIC STAFF MEMBERS" in text
    assert "PARTICULARS OF CLAIMANT" in text
    assert "PARTICULARS OF CLAIM" in text
    assert "Claimant’s Signature" in text
    assert "VERY IMPORTANT" in text


def test_maria_claim_v2_renders_expected_content(monkeypatch, tmp_path):
    dev_reset()
    prepare_manual_templates_for_render(validate=True)
    import_master_data(create_maria_pilot_workbook())

    def _tmp_generated_v2_directory(year: int, month: int, staff_number: str) -> Path:
        return tmp_path / "generated_v2" / str(year) / f"{month:02d}" / str(staff_number)

    monkeypatch.setattr("app_docxtpl.render_claim_v2.generated_v2_directory", _tmp_generated_v2_directory)
    output = render_claim_v2(1008977, 2026, 4)
    text = _doc_text(output)

    assert output == tmp_path / "generated_v2/2026/04/1008977/claim_form_v2_1008977_2026_04.docx"
    assert output.exists()
    assert "Maria Matias" in text
    assert "1008977" in text
    assert "94" in text
    assert "ICT521S & CUS411S" in text
    assert "CUS HORTICULTURE" in text
    assert "CUS GROUP A" in text
    assert "ICT GREY" in text
    assert "ICT Distance" in text
    assert "{{ lecturer_name }}" not in text
    assert "{{" not in text
    assert "}}" not in text
    assert "Demo Clean Lecturer" not in text
    assert "Dummy Qualification" not in text


def test_v2_template_paths_are_separate_from_old_golden_templates():
    assert CLAIM_TEMPLATE_V2 == Path("data/docx_templates_v2/claim_template_v2.docx")
    assert REGISTER_TEMPLATE_V2 == Path("data/docx_templates_v2/attendance_register_template_v2.docx")
