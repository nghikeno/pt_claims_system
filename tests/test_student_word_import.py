from pathlib import Path

from docx import Document

from app.student_word_import import parse_attendance_docx


def make_sample_attendance_docx(path: Path) -> Path:
    doc = Document()
    doc.add_paragraph("CLASS ATTENDANCE SHEET")
    details = doc.add_table(rows=5, cols=2)
    details.rows[0].cells[0].text = "FACULTY"
    details.rows[0].cells[1].text = "Computing and Informatics"
    details.rows[1].cells[0].text = "DEPARTMENT"
    details.rows[1].cells[1].text = "Informatics"
    details.rows[2].cells[0].text = "COURSE NAME"
    details.rows[2].cells[1].text = "Computer User Skills"
    details.rows[3].cells[0].text = "COURSE CODE"
    details.rows[3].cells[1].text = "CUS411S"
    details.rows[4].cells[0].text = "GROUP"
    details.rows[4].cells[1].text = "2"
    table = doc.add_table(rows=6, cols=3)
    table.rows[0].cells[0].text = "NR."
    table.rows[0].cells[1].text = "STUDENT SURNAME & INITIAL"
    table.rows[0].cells[2].text = "STD NR"
    table.rows[1].cells[0].text = "1"
    table.rows[1].cells[1].text = "Haukongo, JL"
    table.rows[1].cells[2].text = "226173453"
    table.rows[2].cells[0].text = "2"
    table.rows[2].cells[1].text = "Venasius, FPN"
    table.rows[2].cells[2].text = "2261755170"
    table.rows[3].cells[0].text = "3"
    table.rows[3].cells[1].text = "Nalukaku, WE"
    table.rows[3].cells[2].text = "225055481"
    table.rows[4].cells[0].text = "NAME OF LECTURER"
    table.rows[4].cells[1].text = "Lonia Nghitotelwa"
    table.rows[4].cells[2].text = "100718"
    table.rows[5].cells[0].text = "SIGNATURE"
    doc.save(path)
    return path


def make_split_column_attendance_docx(path: Path) -> Path:
    doc = Document()
    doc.add_paragraph("CLASS ATTENDANCE SHEET")
    table = doc.add_table(rows=5, cols=6)
    table.rows[0].cells[0].text = "NR."
    table.rows[0].cells[1].text = "STUDENT SURNAME & INITIAL"
    table.rows[0].cells[3].text = "STD NR"
    table.rows[0].cells[4].text = "02-02-26"
    table.rows[1].cells[0].text = "1."
    table.rows[1].cells[1].text = "Haukongo"
    table.rows[1].cells[2].text = "JL"
    table.rows[1].cells[3].text = "226173453"
    table.rows[2].cells[0].text = "2."
    table.rows[2].cells[1].text = "Venasius"
    table.rows[2].cells[2].text = "FPN"
    table.rows[2].cells[3].text = "2261755170"
    table.rows[3].cells[0].text = "3."
    table.rows[3].cells[1].text = "Nalukaku"
    table.rows[3].cells[2].text = "WE"
    table.rows[3].cells[3].text = "225055481"
    table.rows[4].cells[0].text = "SIGNATURE"
    doc.save(path)
    return path


def test_word_parser_extracts_header_and_students(tmp_path):
    path = make_sample_attendance_docx(tmp_path / "attendance.docx")

    parsed = parse_attendance_docx(path)

    assert parsed.header["course_code"] == "CUS411S"
    assert parsed.header["group_label"] == "2"
    assert parsed.students[0] == {
        "student_number": "226173453",
        "surname": "Haukongo",
        "initials": "JL",
        "full_name": "Haukongo JL",
    }
    assert parsed.students[1]["student_number"] == "2261755170"
    assert all(student["surname"] != "NAME OF LECTURER" for student in parsed.students)


def test_word_parser_handles_student_numbers_as_text(tmp_path):
    path = make_sample_attendance_docx(tmp_path / "attendance.docx")
    parsed = parse_attendance_docx(path)

    assert isinstance(parsed.students[0]["student_number"], str)


def test_word_parser_handles_split_surname_initial_columns_without_row_number_in_name(tmp_path):
    path = make_split_column_attendance_docx(tmp_path / "split.docx")

    parsed = parse_attendance_docx(path)

    assert parsed.students[0]["student_number"] == "226173453"
    assert parsed.students[0]["surname"] == "Haukongo"
    assert parsed.students[0]["initials"] == "JL"
    assert parsed.students[0]["full_name"] == "Haukongo JL"
    assert parsed.students[1]["student_number"] == "2261755170"
    assert parsed.students[1]["surname"] == "Venasius"
    assert parsed.students[1]["initials"] == "FPN"
    for student in parsed.students:
        assert "1." not in student["surname"]
        assert "1." not in student["initials"]
        assert "1." not in student["full_name"]
        assert "2." not in student["surname"]
        assert "2." not in student["initials"]
        assert "2." not in student["full_name"]


def test_word_parser_detects_bank_text_but_does_not_extract_bank_rows(tmp_path):
    path = make_sample_attendance_docx(tmp_path / "attendance.docx")
    doc = Document(path)
    doc.add_paragraph("Bank account number: 123456")
    doc.save(path)

    parsed = parse_attendance_docx(path)

    assert "Bank details detected and ignored." in parsed.warnings
