from __future__ import annotations

from pathlib import Path

from docx import Document

from app.create_maria_pilot_workbook import EXPECTED_AMOUNT, EXPECTED_HOURS, EXPECTED_SESSIONS, create_maria_pilot_workbook
from app.dev_reset import dev_reset
from app.import_master_data import import_master_data
from app.session_generator import generate_monthly_sessions
from app_docxtpl.manual_templates import (
    MANUAL_CLAIM_TEMPLATE_V2,
    MANUAL_REGISTER_TEMPLATE_V2,
    prepare_manual_templates_for_render,
    template_contains_text,
)
from app_docxtpl.render_documents_v2 import render_documents_v2


UNRESOLVED_MARKERS = ("{{", "}}", "{%tr", "{% tr", "claim_rows", "lecturer_name", "row.", "student.", "session_1_date")


def _doc_text(path: Path) -> str:
    document = Document(path)
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


def test_manual_templates_exist_and_hashes_are_unchanged_after_copy():
    assert MANUAL_CLAIM_TEMPLATE_V2.exists()
    assert MANUAL_REGISTER_TEMPLATE_V2.exists()

    info = prepare_manual_templates_for_render(validate=True)

    assert info["before"]["claim"] == info["after"]["claim"]
    assert info["before"]["register"] == info["after"]["register"]
    assert info["render_paths"]["claim"].exists()
    assert info["render_paths"]["register"].exists()
    assert info["manual_claim_hash"] == info["render_claim_hash"]
    assert info["manual_register_hash"] == info["render_register_hash"]


def _patch_v2_output_dirs(monkeypatch, tmp_path):
    def _tmp_generated_v2_directory(year: int, month: int, staff_number: str) -> Path:
        return tmp_path / "generated_v2" / str(year) / f"{month:02d}" / str(staff_number)

    monkeypatch.setattr("app_docxtpl.render_documents_v2.generated_v2_directory", _tmp_generated_v2_directory)
    monkeypatch.setattr("app_docxtpl.render_claim_v2.generated_v2_directory", _tmp_generated_v2_directory)
    monkeypatch.setattr("app_docxtpl.render_register_v2.generated_v2_directory", _tmp_generated_v2_directory)


def test_rendering_maria_from_manual_templates_succeeds_and_preserves_totals(monkeypatch, tmp_path):
    dev_reset()
    import_master_data(create_maria_pilot_workbook())

    _patch_v2_output_dirs(monkeypatch, tmp_path)
    result = render_documents_v2(1008977, 2026, 4)
    sessions_df = generate_monthly_sessions(1008977, 2026, 4)

    assert result["claim_path"].exists()
    assert result["register_paths"]
    assert all(path.exists() for path in result["register_paths"])
    assert len(sessions_df) == EXPECTED_SESSIONS
    assert round(float(sessions_df["hours"].sum()), 2) == EXPECTED_HOURS
    assert round(float(sessions_df["amount"].sum()), 2) == EXPECTED_AMOUNT
    assert result["total_sessions"] == EXPECTED_SESSIONS
    assert result["total_hours"] == "94"
    assert result["total_amount"] == EXPECTED_AMOUNT
    assert result["template_info"]["before"] == result["template_info"]["after"]
    assert result["provenance_path"]
    assert Path(result["provenance_path"]).exists()


def test_render_provenance_contains_manual_template_paths_and_hashes(monkeypatch, tmp_path):
    dev_reset()
    import_master_data(create_maria_pilot_workbook())

    _patch_v2_output_dirs(monkeypatch, tmp_path)
    result = render_documents_v2(1008977, 2026, 4)
    provenance = Path(result["provenance_path"]).read_text(encoding="utf-8")

    assert str(MANUAL_CLAIM_TEMPLATE_V2) in provenance
    assert str(MANUAL_REGISTER_TEMPLATE_V2) in provenance
    assert result["template_info"]["manual_claim_hash"] in provenance
    assert result["template_info"]["manual_register_hash"] in provenance
    assert "number of register files:" in provenance


def test_render_documents_v2_deletes_stale_output_before_rendering(monkeypatch, tmp_path):
    dev_reset()
    import_master_data(create_maria_pilot_workbook())

    _patch_v2_output_dirs(monkeypatch, tmp_path)
    first = render_documents_v2(1008977, 2026, 4)
    stale_file = Path(first["output_dir"]) / "stale-output-marker.txt"
    stale_file.write_text("stale", encoding="utf-8")
    assert stale_file.exists()

    second = render_documents_v2(1008977, 2026, 4)

    assert not stale_file.exists()
    assert Path(second["provenance_path"]).exists()


def test_template_contains_text_detects_known_docx_marker():
    assert template_contains_text(MANUAL_CLAIM_TEMPLATE_V2, "CLAIM")
    assert not template_contains_text(MANUAL_CLAIM_TEMPLATE_V2, "THIS MARKER SHOULD NOT EXIST")


def test_manual_template_rendered_maria_claim_has_values_and_no_unresolved_placeholders(monkeypatch, tmp_path):
    dev_reset()
    import_master_data(create_maria_pilot_workbook())

    _patch_v2_output_dirs(monkeypatch, tmp_path)
    result = render_documents_v2(1008977, 2026, 4)
    text = _doc_text(result["claim_path"])

    for expected in (
        "Maria Matias",
        "1008977",
        "Master of Science in Information Technology",
        "0183-0102",
        "ICT521S & CUS411S",
        "94",
        "CUS HORTICULTURE",
        "CUS GROUP A",
        "CUS GROUP B",
        "ICT GREY",
        "ICT Distance",
        "Claims older than 3 months",
    ):
        assert expected in text
    for forbidden in UNRESOLVED_MARKERS + ("Demo Clean Lecturer", "Dummy Qualification"):
        assert forbidden not in text


def test_manual_template_rendered_maria_registers_have_values_and_no_unresolved_placeholders(monkeypatch, tmp_path):
    dev_reset()
    import_master_data(create_maria_pilot_workbook())

    _patch_v2_output_dirs(monkeypatch, tmp_path)
    result = render_documents_v2(1008977, 2026, 4)
    combined_text = "\n".join(_doc_text(path) for path in result["register_paths"])

    assert Path(result["output_dir"], "registers").exists()
    for expected in ("Maria Matias", "1008977", "07-04-26", "11:30-12:30", "PilotSurname001"):
        assert expected in combined_text
    for forbidden in UNRESOLVED_MARKERS + ("Demo Clean Lecturer", "Amunyela"):
        assert forbidden not in combined_text

    group_names = ["CUS HORTICULTURE", "CUS GROUP A", "CUS GROUP B", "ICT GREY", "ICT Distance"]
    for path in result["register_paths"]:
        text = _doc_text(path)
        groups_present = [group for group in group_names if group in text]
        assert len(groups_present) <= 1
