from pathlib import Path

from docx import Document
import pytest

from app.claim_form_generator import generate_claim_form
from app.seed_data import seed_database
from app.session_generator import generate_monthly_sessions


@pytest.fixture(autouse=True)
def seeded_database():
    seed_database()


def _doc_text(path: Path) -> str:
    document = Document(path)
    paragraphs = [paragraph.text for paragraph in document.paragraphs]
    table_text = []
    for table in document.tables:
        for row in table.rows:
            table_text.extend(cell.text for cell in row.cells)
    return "\n".join(paragraphs + table_text)


def test_claim_form_docx_file_is_created(tmp_path):
    sessions_df = generate_monthly_sessions(1, 2026, 2)
    output_path = tmp_path / "claim.docx"

    generate_claim_form(sessions_df, output_path, 2026, 2)

    assert output_path.exists()


def test_claim_form_contains_required_details_and_total_hours(tmp_path):
    sessions_df = generate_monthly_sessions(1, 2026, 2)
    output_path = tmp_path / "claim.docx"

    generate_claim_form(sessions_df, output_path, 2026, 2)
    text = _doc_text(output_path)

    assert "Lonia Nghitotelwa" in text
    assert "100718" in text
    assert "0183-0102" in text
    assert "410.00" in text
    assert "68" in text
    assert "Field\nValue" not in text[:500]
    assert "PARTICULARS OF CLAIMANT" in text
