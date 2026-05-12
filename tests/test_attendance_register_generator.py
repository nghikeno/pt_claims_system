from pathlib import Path

from docx import Document
import pytest

from app.attendance_register_generator import generate_attendance_register_pack
from app.database import get_connection
from app.seed_data import seed_database
from app.session_generator import generate_monthly_sessions


@pytest.fixture(autouse=True)
def seeded_database():
    seed_database()


def _doc_text(path: Path) -> str:
    document = Document(path)
    paragraphs = [paragraph.text for paragraph in document.paragraphs]
    table_text = []
    for table in document.tables:
        for row in table.rows:
            table_text.extend(cell.text for cell in row.cells)
    return "\n".join(paragraphs + table_text)


def test_attendance_register_docx_file_is_created(tmp_path):
    sessions_df = generate_monthly_sessions(1, 2026, 2)
    output_path = tmp_path / "attendance.docx"

    generate_attendance_register_pack(sessions_df, output_path, 2026, 2)

    assert output_path.exists()


def test_attendance_register_contains_required_details(tmp_path):
    sessions_df = generate_monthly_sessions(1, 2026, 2)
    output_path = tmp_path / "attendance.docx"

    generate_attendance_register_pack(sessions_df, output_path, 2026, 2)
    text = _doc_text(output_path)

    assert "Lonia Nghitotelwa" in text
    assert "100718" in text
    assert "CUS411S" in text
    assert "Group 1" in text


def test_groups_with_more_than_five_sessions_are_split(tmp_path):
    with get_connection() as conn:
        lecturer_id = conn.execute("SELECT id FROM lecturers WHERE staff_number = '100718'").fetchone()["id"]
        group_id = conn.execute("SELECT id FROM student_groups WHERE group_name = 'Group 7'").fetchone()["id"]
        for day, start, end in [
            ("Tuesday", "12:00", "13:00"),
            ("Wednesday", "08:00", "09:00"),
            ("Thursday", "12:00", "13:00"),
        ]:
            conn.execute(
                """
                INSERT INTO timetable_entries (
                    lecturer_id, group_id, day_of_week, start_time, end_time,
                    effective_start_date, effective_end_date, active
                )
                VALUES (?, ?, ?, ?, ?, '2026-02-01', '2026-02-28', 1)
                """,
                (lecturer_id, group_id, day, start, end),
            )
    sessions_df = generate_monthly_sessions(1, 2026, 2)
    output_path = tmp_path / "attendance_split.docx"

    generate_attendance_register_pack(sessions_df, output_path, 2026, 2)
    text = _doc_text(output_path)

    assert "REG-100718-CUS411S-GROUP-7-2026-02-P2" in text
    assert "Page 2 of" in text


def test_no_students_creates_blank_numbered_rows(tmp_path):
    sessions_df = generate_monthly_sessions(1, 2026, 2)
    output_path = tmp_path / "attendance_blank_rows.docx"

    generate_attendance_register_pack(sessions_df, output_path, 2026, 2)
    document = Document(output_path)
    register_tables = [
        table
        for table in document.tables
        if table.rows and table.rows[0].cells[0].text == "NR."
    ]

    assert register_tables
    assert len(register_tables[0].rows) == 32
    assert register_tables[0].rows[-1].cells[0].text == "30"
